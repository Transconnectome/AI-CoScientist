#!/usr/bin/env python3
"""Generate a professional DOCX report from abstract improvement files."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import json

# Read all files
with open('abstract_original.txt', 'r', encoding='utf-8') as f:
    original_text = f.read()

with open('abstract_improved_v2.txt', 'r', encoding='utf-8') as f:
    improved_text = f.read()

with open('abstract_evaluation.json', 'r', encoding='utf-8') as f:
    original_eval = json.load(f)

with open('abstract_evaluation_v2.json', 'r', encoding='utf-8') as f:
    improved_eval = json.load(f)

# Create document
doc = Document()

# Title
title = doc.add_heading('Abstract Improvement Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle with date
subtitle = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(11)
subtitle_run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()  # Space

# Executive Summary
doc.add_heading('Executive Summary', 1)
summary = doc.add_paragraph()
summary.add_run('Overall Score: ').bold = True
summary.add_run(f'{original_eval["overall"]["score"]:.2f} → {improved_eval["overall"]["score"]:.2f} ')
improvement_run = summary.add_run(f'(+{improved_eval["overall"]["score"] - original_eval["overall"]["score"]:.2f})')
improvement_run.bold = True

word_count_para = doc.add_paragraph()
word_count_para.add_run('Word Count: ').bold = True
word_count_para.add_run(f'230 → 170 (-60 words, -26.1%)')

doc.add_paragraph()

# Score Comparison Table
doc.add_heading('Score Comparison', 1)
table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
header_cells[0].text = 'Dimension'
header_cells[1].text = 'Original'
header_cells[2].text = 'Improved'
header_cells[3].text = 'Change'

# Make header bold
for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Data rows
dimensions = [
    ('Overall', 'overall'),
    ('Clarity', 'clarity'),
    ('Completeness', 'completeness'),
    ('Conciseness', 'conciseness'),
    ('Impact', 'impact')
]

for idx, (name, key) in enumerate(dimensions, 1):
    row = table.rows[idx]
    orig_score = original_eval[key]['score'] if key in original_eval else original_eval[key]
    imp_score = improved_eval[key]['score'] if key in improved_eval else improved_eval[key]
    change = imp_score - orig_score

    row.cells[0].text = name
    row.cells[1].text = f'{orig_score:.2f}'
    row.cells[2].text = f'{imp_score:.2f}'
    row.cells[3].text = f'{change:+.2f}'

doc.add_page_break()

# Original Abstract
doc.add_heading('Original Abstract', 1)
doc.add_paragraph(f'Word Count: 230 | Score: {original_eval["overall"]["score"]:.2f}/10')
original_para = doc.add_paragraph(original_text)
original_para.style = 'Quote'

doc.add_paragraph()

# Improved Abstract
doc.add_heading('Improved Abstract (Version 2)', 1)
doc.add_paragraph(f'Word Count: 170 | Score: {improved_eval["overall"]["score"]:.2f}/10')
improved_para = doc.add_paragraph(improved_text)
improved_para.style = 'Intense Quote'
for run in improved_para.runs:
    run.bold = True

doc.add_page_break()

# Detailed Evaluation
doc.add_heading('Detailed Evaluation of Improved Version', 1)

# Dimensional Scores
doc.add_heading('Dimensional Scores', 2)

for dim_name, dim_key in [('Clarity', 'clarity'), ('Completeness', 'completeness'),
                          ('Conciseness', 'conciseness'), ('Impact', 'impact')]:
    dim_data = improved_eval[dim_key]

    p = doc.add_paragraph()
    p.add_run(f'{dim_name}: ').bold = True
    p.add_run(f'{dim_data["score"]:.2f}/10')

    justification = doc.add_paragraph(dim_data['justification'])
    justification.paragraph_format.left_indent = Inches(0.3)
    justification.paragraph_format.space_after = Pt(6)

doc.add_paragraph()

# Strengths
doc.add_heading('Strengths', 2)
strengths = improved_eval.get('strengths', [])
for strength in strengths:
    p = doc.add_paragraph(strength, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

doc.add_paragraph()

# Areas for Improvement
doc.add_heading('Areas for Improvement', 2)
weaknesses = improved_eval.get('weaknesses', [])
for weakness in weaknesses:
    p = doc.add_paragraph(weakness, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

doc.add_paragraph()

# Improvement Suggestions
doc.add_heading('Further Improvement Suggestions', 2)
suggestions = improved_eval.get('suggestions', [])
for idx, suggestion in enumerate(suggestions, 1):
    p = doc.add_paragraph(suggestion, style='List Number')
    p.paragraph_format.space_after = Pt(3)

# Save
output_file = 'abstract_improvement_report.docx'
doc.save(output_file)

print('=' * 80)
print('📄 DOCX REPORT GENERATED')
print('=' * 80)
print(f'\nFile: {output_file}')
print(f'Size: {len(original_text.split())} → {len(improved_text.split())} words')
print(f'Score: {original_eval["overall"]["score"]:.2f} → {improved_eval["overall"]["score"]:.2f}')
print('\n✅ Report saved successfully!')
print('\nReport includes:')
print('  • Executive Summary')
print('  • Score Comparison Table')
print('  • Original Abstract')
print('  • Improved Abstract (highlighted)')
print('  • Detailed Evaluation')
print('  • Strengths & Weaknesses')
print('  • Improvement Suggestions')
