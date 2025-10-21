#!/usr/bin/env python3
"""Complete abstract improvement workflow with DOCX report generation."""

import sys
sys.path.insert(0, 'scripts')

from anthropic import Anthropic
from section_evaluator import SectionEvaluator, SectionType, EvaluationContext
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import json

print('='*80)
print('COMPLETE ABSTRACT IMPROVEMENT WORKFLOW')
print('='*80)

# Load API key
with open('.env', 'r') as f:
    for line in f:
        if line.startswith('ANTHROPIC_API_KEY='):
            api_key = line.strip().split('=', 1)[1]
            break

# Read original abstract
with open('abstract_original.txt', 'r') as f:
    original = f.read()

print(f'\n📄 Original Abstract: {len(original.split())} words')

# Step 1: Evaluate original
print('\n🔍 Step 1: Evaluating original abstract...')
evaluator = SectionEvaluator()
original_eval = evaluator.evaluate_section(
    SectionType.ABSTRACT,
    original,
    EvaluationContext.STANDALONE
)
print(f'   Original Score: {original_eval["overall"]["score"]:.2f}/10')

# Save original evaluation
with open('abstract_evaluation.json', 'w') as f:
    json.dump(original_eval, f, indent=2, ensure_ascii=False)

# Step 2: Improve with LLM
print('\n🔧 Step 2: Generating improved version with Claude AI...')

improvement_prompt = f'''Scientific writing expert: Improve this abstract.

ORIGINAL (words: {len(original.split())}, score: {original_eval["overall"]["score"]:.2f}/10):
{original}

EVALUATION FEEDBACK:
- Clarity: {original_eval["clarity"]["score"]:.2f}/10 - Technical jargon limits accessibility
- Completeness: {original_eval["completeness"]["score"]:.2f}/10 - Good but missing age range
- Conciseness: {original_eval["conciseness"]["score"]:.2f}/10 - Too verbose
- Impact: {original_eval["impact"]["score"]:.2f}/10 - Strong findings, weak presentation

IMPROVEMENTS NEEDED:
1. Lead with paradoxical finding as hook (most striking result first)
2. Add ONE quantitative result to anchor findings
3. Simplify technical terms with brief context
4. Add children's age range
5. Reduce repetition of key terms
6. Focus on primary outcome
7. Target 170-190 words

Provide ONLY the improved abstract.'''

client = Anthropic(api_key=api_key)
response = client.messages.create(
    model='claude-sonnet-4-5-20250929',
    max_tokens=2048,
    temperature=0.3,
    messages=[{'role': 'user', 'content': improvement_prompt}]
)

improved = response.content[0].text.strip()
print(f'   Improved: {len(improved.split())} words')

# Save improved version
with open('abstract_improved.txt', 'w') as f:
    f.write(improved)

# Step 3: Evaluate improved
print('\n🔍 Step 3: Evaluating improved abstract...')
improved_eval = evaluator.evaluate_section(
    SectionType.ABSTRACT,
    improved,
    EvaluationContext.STANDALONE
)
print(f'   Improved Score: {improved_eval["overall"]["score"]:.2f}/10')

# Save improved evaluation
with open('abstract_evaluation_improved.json', 'w') as f:
    json.dump(improved_eval, f, indent=2, ensure_ascii=False)

# Step 4: Generate DOCX report
print('\n📊 Step 4: Generating DOCX report...')

doc = Document()

# Title
title = doc.add_heading('Abstract Improvement Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(11)
subtitle_run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# Executive Summary
doc.add_heading('Executive Summary', 1)
summary = doc.add_paragraph()
summary.add_run('Overall Score: ').bold = True
summary.add_run(f'{original_eval["overall"]["score"]:.2f} → {improved_eval["overall"]["score"]:.2f} ')
improvement_run = summary.add_run(f'(+{improved_eval["overall"]["score"] - original_eval["overall"]["score"]:.2f})')
improvement_run.bold = True

word_para = doc.add_paragraph()
word_para.add_run('Word Count: ').bold = True
word_para.add_run(f'{len(original.split())} → {len(improved.split())} (-{len(original.split()) - len(improved.split())} words)')

doc.add_paragraph()

# Score Comparison Table
doc.add_heading('Score Comparison', 1)
table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'

header_cells = table.rows[0].cells
header_cells[0].text = 'Dimension'
header_cells[1].text = 'Original'
header_cells[2].text = 'Improved'
header_cells[3].text = 'Change'

for cell in header_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

dimensions = [
    ('Overall', 'overall'),
    ('Clarity', 'clarity'),
    ('Completeness', 'completeness'),
    ('Conciseness', 'conciseness'),
    ('Impact', 'impact')
]

for idx, (name, key) in enumerate(dimensions, 1):
    row = table.rows[idx]
    orig_score = original_eval[key]['score'] if key in original_eval else original_eval[key]
    imp_score = improved_eval[key]['score'] if key in improved_eval else improved_eval[key]
    change = imp_score - orig_score

    row.cells[0].text = name
    row.cells[1].text = f'{orig_score:.2f}'
    row.cells[2].text = f'{imp_score:.2f}'
    row.cells[3].text = f'{change:+.2f}'

doc.add_page_break()

# Original Abstract
doc.add_heading('Original Abstract', 1)
doc.add_paragraph(f'Word Count: {len(original.split())} | Score: {original_eval["overall"]["score"]:.2f}/10')
original_para = doc.add_paragraph(original)
original_para.style = 'Quote'

doc.add_paragraph()

# Improved Abstract
doc.add_heading('Improved Abstract', 1)
doc.add_paragraph(f'Word Count: {len(improved.split())} | Score: {improved_eval["overall"]["score"]:.2f}/10')
improved_para = doc.add_paragraph(improved)
improved_para.style = 'Intense Quote'
for run in improved_para.runs:
    run.bold = True

doc.add_page_break()

# Detailed Evaluation
doc.add_heading('Detailed Evaluation', 1)

for dim_name, dim_key in [('Clarity', 'clarity'), ('Completeness', 'completeness'),
                          ('Conciseness', 'conciseness'), ('Impact', 'impact')]:
    dim_data = improved_eval[dim_key]
    p = doc.add_paragraph()
    p.add_run(f'{dim_name}: ').bold = True
    p.add_run(f'{dim_data["score"]:.2f}/10')
    justification = doc.add_paragraph(dim_data['justification'])
    justification.paragraph_format.left_indent = Inches(0.3)
    justification.paragraph_format.space_after = Pt(6)

doc.add_paragraph()

# Strengths
doc.add_heading('Strengths', 2)
for strength in improved_eval.get('strengths', []):
    p = doc.add_paragraph(strength, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

doc.add_paragraph()

# Weaknesses
doc.add_heading('Areas for Improvement', 2)
for weakness in improved_eval.get('weaknesses', []):
    p = doc.add_paragraph(weakness, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

doc.add_paragraph()

# Suggestions
doc.add_heading('Further Improvement Suggestions', 2)
for idx, suggestion in enumerate(improved_eval.get('suggestions', []), 1):
    p = doc.add_paragraph(suggestion, style='List Number')
    p.paragraph_format.space_after = Pt(3)

# Save DOCX
output_file = 'abstract_improvement_report.docx'
doc.save(output_file)

print(f'   ✅ {output_file}')

# Final Summary
print('\n' + '='*80)
print('RESULTS SUMMARY')
print('='*80)
print(f'\n📊 Scores:')
print(f'   Overall:      {original_eval["overall"]["score"]:.2f} → {improved_eval["overall"]["score"]:.2f} ({improved_eval["overall"]["score"] - original_eval["overall"]["score"]:+.2f})')
print(f'   Clarity:      {original_eval["clarity"]["score"]:.2f} → {improved_eval["clarity"]["score"]:.2f} ({improved_eval["clarity"]["score"] - original_eval["clarity"]["score"]:+.2f})')
print(f'   Completeness: {original_eval["completeness"]["score"]:.2f} → {improved_eval["completeness"]["score"]:.2f} ({improved_eval["completeness"]["score"] - original_eval["completeness"]["score"]:+.2f})')
print(f'   Conciseness:  {original_eval["conciseness"]["score"]:.2f} → {improved_eval["conciseness"]["score"]:.2f} ({improved_eval["conciseness"]["score"] - original_eval["conciseness"]["score"]:+.2f})')
print(f'   Impact:       {original_eval["impact"]["score"]:.2f} → {improved_eval["impact"]["score"]:.2f} ({improved_eval["impact"]["score"] - original_eval["impact"]["score"]:+.2f})')

print(f'\n📝 Word Count: {len(original.split())} → {len(improved.split())} ({len(improved.split()) - len(original.split())} words)')

print('\n📁 Generated Files:')
print('   • abstract_improved.txt')
print('   • abstract_evaluation.json')
print('   • abstract_evaluation_improved.json')
print('   • abstract_improvement_report.docx ⭐')

print('\n✅ Complete!')
