#!/usr/bin/env python3
"""Iterative abstract improvement with automatic feedback integration."""

import sys
sys.path.insert(0, 'scripts')

from anthropic import Anthropic
from section_evaluator import SectionEvaluator, SectionType, EvaluationContext
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import json

print('='*80)
print('ITERATIVE ABSTRACT IMPROVEMENT WITH AUTO-FEEDBACK')
print('='*80)

# Configuration
TARGET_SCORE = 8.5
MAX_ITERATIONS = 5
MIN_IMPROVEMENT = 0.05  # Minimum improvement to continue

# Load API key
with open('.env', 'r') as f:
    for line in f:
        if line.startswith('ANTHROPIC_API_KEY='):
            api_key = line.strip().split('=', 1)[1]
            break

# Read original abstract
with open('abstract_original.txt', 'r') as f:
    original_text = f.read()

print(f'\n📄 Original Abstract: {len(original_text.split())} words')

# Initialize
evaluator = SectionEvaluator()
client = Anthropic(api_key=api_key)

# Track improvement history
history = []
current_text = original_text

# Initial evaluation
print('\n🔍 Initial Evaluation...')
initial_eval = evaluator.evaluate_section(
    SectionType.ABSTRACT,
    current_text,
    EvaluationContext.STANDALONE
)
print(f'   Initial Score: {initial_eval["overall"]["score"]:.2f}/10')

history.append({
    'iteration': 0,
    'text': current_text,
    'evaluation': initial_eval,
    'word_count': len(current_text.split())
})

# Iterative improvement loop
for iteration in range(1, MAX_ITERATIONS + 1):
    print(f'\n{"="*80}')
    print(f'ITERATION {iteration}')
    print(f'{"="*80}')

    current_eval = history[-1]['evaluation']
    current_score = current_eval['overall']['score']

    # Check if target reached
    if current_score >= TARGET_SCORE:
        print(f'✅ Target score {TARGET_SCORE} reached! ({current_score:.2f})')
        break

    # Identify weak dimensions (score < 8.0)
    weak_dimensions = []
    for dim in ['clarity', 'completeness', 'conciseness', 'impact']:
        if current_eval[dim]['score'] < 8.0:
            weak_dimensions.append({
                'name': dim,
                'score': current_eval[dim]['score'],
                'justification': current_eval[dim]['justification']
            })

    # Sort by score (weakest first)
    weak_dimensions.sort(key=lambda x: x['score'])

    print(f'\n🎯 Current Score: {current_score:.2f}/10')
    print(f'📊 Weak Dimensions:')
    for dim in weak_dimensions:
        print(f'   • {dim["name"].title()}: {dim["score"]:.2f}/10')

    # Build improvement prompt with automatic feedback integration
    improvement_prompt = f'''Scientific writing expert: Improve this abstract to achieve a score of 8.5+.

CURRENT VERSION (words: {len(current_text.split())}, score: {current_score:.2f}/10):
{current_text}

CURRENT EVALUATION:
Overall Score: {current_score:.2f}/10

Dimensional Scores:
'''

    for dim in ['clarity', 'completeness', 'conciseness', 'impact']:
        score = current_eval[dim]['score']
        improvement_prompt += f'- {dim.title()}: {score:.2f}/10'
        if score < 8.0:
            improvement_prompt += f' ⚠️ NEEDS IMPROVEMENT'
        improvement_prompt += '\n'

    # Add specific weaknesses
    if current_eval.get('weaknesses'):
        improvement_prompt += f'\nKEY WEAKNESSES TO ADDRESS:\n'
        for i, weakness in enumerate(current_eval['weaknesses'][:5], 1):
            improvement_prompt += f'{i}. {weakness}\n'

    # Add specific suggestions (AUTOMATIC FEEDBACK INTEGRATION)
    if current_eval.get('suggestions'):
        improvement_prompt += f'\nSPECIFIC IMPROVEMENT SUGGESTIONS:\n'
        for i, suggestion in enumerate(current_eval['suggestions'][:6], 1):
            improvement_prompt += f'{i}. {suggestion}\n'

    # Add focus areas based on weak dimensions
    if weak_dimensions:
        improvement_prompt += f'\nPRIORITY FOCUS AREAS:\n'
        for i, dim in enumerate(weak_dimensions[:3], 1):
            improvement_prompt += f'{i}. **{dim["name"].upper()}** (current: {dim["score"]:.2f}/10)\n'
            improvement_prompt += f'   Issue: {dim["justification"][:200]}...\n'

    improvement_prompt += f'''
TARGET: Achieve overall score of {TARGET_SCORE}+ while maintaining strengths.

CONSTRAINTS:
- Word count: 160-180 words (currently {len(current_text.split())} words)
- Maintain completeness and impact
- Improve clarity and conciseness
- Address ALL suggestions above

Provide ONLY the improved abstract text.'''

    # Generate improvement
    print(f'\n🔧 Generating Improvement (Iteration {iteration})...')
    response = client.messages.create(
        model='claude-sonnet-4-5-20250929',
        max_tokens=2048,
        temperature=0.3,
        messages=[{'role': 'user', 'content': improvement_prompt}]
    )

    improved_text = response.content[0].text.strip()

    # Evaluate improvement
    print(f'📊 Evaluating Improvement...')
    improved_eval = evaluator.evaluate_section(
        SectionType.ABSTRACT,
        improved_text,
        EvaluationContext.STANDALONE
    )

    improved_score = improved_eval['overall']['score']
    score_change = improved_score - current_score

    print(f'\n📈 Results:')
    print(f'   Previous: {current_score:.2f}/10')
    print(f'   Current:  {improved_score:.2f}/10 ({score_change:+.2f})')
    print(f'   Words: {len(current_text.split())} → {len(improved_text.split())} ({len(improved_text.split()) - len(current_text.split()):+d})')

    # Record history
    history.append({
        'iteration': iteration,
        'text': improved_text,
        'evaluation': improved_eval,
        'word_count': len(improved_text.split()),
        'score_change': score_change
    })

    # Check for meaningful improvement
    if score_change < MIN_IMPROVEMENT and iteration > 1:
        print(f'\n⚠️  Minimal improvement ({score_change:.2f}). Stopping iteration.')
        break

    # Update current text for next iteration
    current_text = improved_text

# Find best version
best_version = max(history, key=lambda x: x['evaluation']['overall']['score'])
best_iteration = best_version['iteration']
final_score = best_version['evaluation']['overall']['score']

print(f'\n{"="*80}')
print(f'IMPROVEMENT COMPLETE')
print(f'{"="*80}')
print(f'\n🏆 Best Version: Iteration {best_iteration}')
print(f'📊 Score Improvement: {initial_eval["overall"]["score"]:.2f} → {final_score:.2f} (+{final_score - initial_eval["overall"]["score"]:.2f})')
print(f'📝 Word Count: {len(original_text.split())} → {best_version["word_count"]} ({best_version["word_count"] - len(original_text.split()):+d})')

# Save best version
with open('abstract_improved_iterative.txt', 'w') as f:
    f.write(best_version['text'])

with open('abstract_evaluation_iterative.json', 'w') as f:
    json.dump(best_version['evaluation'], f, indent=2, ensure_ascii=False)

# Save improvement history
with open('improvement_history.json', 'w') as f:
    history_export = []
    for h in history:
        history_export.append({
            'iteration': h['iteration'],
            'score': h['evaluation']['overall']['score'],
            'word_count': h['word_count'],
            'scores_by_dimension': {
                dim: h['evaluation'][dim]['score']
                for dim in ['clarity', 'completeness', 'conciseness', 'impact']
            }
        })
    json.dump(history_export, f, indent=2)

print(f'\n✅ Files saved:')
print(f'   • abstract_improved_iterative.txt')
print(f'   • abstract_evaluation_iterative.json')
print(f'   • improvement_history.json')

# Generate comprehensive DOCX report
print(f'\n📊 Generating Comprehensive Report...')

doc = Document()

# Title
title = doc.add_heading('Iterative Abstract Improvement Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(11)
subtitle_run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# Executive Summary
doc.add_heading('Executive Summary', 1)

summary = doc.add_paragraph()
summary.add_run('Improvement Method: ').bold = True
summary.add_run('Iterative refinement with automatic feedback integration')

iterations_para = doc.add_paragraph()
iterations_para.add_run('Total Iterations: ').bold = True
iterations_para.add_run(f'{len(history) - 1} (max {MAX_ITERATIONS})')

score_para = doc.add_paragraph()
score_para.add_run('Score Improvement: ').bold = True
score_para.add_run(f'{initial_eval["overall"]["score"]:.2f} → {final_score:.2f} ')
improvement_run = score_para.add_run(f'(+{final_score - initial_eval["overall"]["score"]:.2f})')
improvement_run.bold = True
improvement_run.font.color.rgb = RGBColor(0, 128, 0)

word_para = doc.add_paragraph()
word_para.add_run('Word Count: ').bold = True
word_para.add_run(f'{len(original_text.split())} → {best_version["word_count"]} ({best_version["word_count"] - len(original_text.split()):+d} words)')

doc.add_paragraph()

# Iteration History
doc.add_heading('Improvement Progression', 1)

history_table = doc.add_table(rows=len(history) + 1, cols=6)
history_table.style = 'Light Grid Accent 1'

# Header
header_cells = history_table.rows[0].cells
headers = ['Iteration', 'Overall', 'Clarity', 'Completeness', 'Conciseness', 'Impact']
for i, header in enumerate(headers):
    header_cells[i].text = header
    for paragraph in header_cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Data rows
for idx, h in enumerate(history, 1):
    row = history_table.rows[idx]
    eval_data = h['evaluation']

    row.cells[0].text = f"V{h['iteration']}"
    row.cells[1].text = f"{eval_data['overall']['score']:.2f}"
    row.cells[2].text = f"{eval_data['clarity']['score']:.2f}"
    row.cells[3].text = f"{eval_data['completeness']['score']:.2f}"
    row.cells[4].text = f"{eval_data['conciseness']['score']:.2f}"
    row.cells[5].text = f"{eval_data['impact']['score']:.2f}"

    # Highlight best version
    if h['iteration'] == best_iteration:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_page_break()

# Original Abstract
doc.add_heading('Original Abstract', 1)
doc.add_paragraph(f'Word Count: {len(original_text.split())} | Score: {initial_eval["overall"]["score"]:.2f}/10')
original_para = doc.add_paragraph(original_text)
original_para.style = 'Quote'

doc.add_paragraph()

# Best Improved Abstract
doc.add_heading(f'Best Improved Abstract (Iteration {best_iteration})', 1)
doc.add_paragraph(f'Word Count: {best_version["word_count"]} | Score: {final_score:.2f}/10')
improved_para = doc.add_paragraph(best_version['text'])
improved_para.style = 'Intense Quote'
for run in improved_para.runs:
    run.bold = True

doc.add_page_break()

# Detailed Evaluation
doc.add_heading('Detailed Evaluation (Best Version)', 1)

best_eval = best_version['evaluation']

for dim_name, dim_key in [('Clarity', 'clarity'), ('Completeness', 'completeness'),
                          ('Conciseness', 'conciseness'), ('Impact', 'impact')]:
    dim_data = best_eval[dim_key]
    p = doc.add_paragraph()
    p.add_run(f'{dim_name}: ').bold = True
    p.add_run(f'{dim_data["score"]:.2f}/10')
    justification = doc.add_paragraph(dim_data['justification'])
    justification.paragraph_format.left_indent = Inches(0.3)
    justification.paragraph_format.space_after = Pt(6)

doc.add_paragraph()

# Strengths
doc.add_heading('Strengths', 2)
for strength in best_eval.get('strengths', []):
    p = doc.add_paragraph(strength, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

doc.add_paragraph()

# Remaining Areas for Improvement
doc.add_heading('Remaining Areas for Improvement', 2)
for weakness in best_eval.get('weaknesses', []):
    p = doc.add_paragraph(weakness, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

# Save report
output_file = 'abstract_iterative_improvement_report.docx'
doc.save(output_file)

print(f'   ✅ {output_file}')

# Final summary
print(f'\n{"="*80}')
print(f'FINAL SUMMARY')
print(f'{"="*80}')
print(f'\n📊 Score Progression:')
for h in history:
    symbol = '🏆' if h['iteration'] == best_iteration else '  '
    print(f'{symbol} V{h["iteration"]}: {h["evaluation"]["overall"]["score"]:.2f}/10 ({h["word_count"]} words)')

print(f'\n📈 Dimensional Improvements:')
for dim in ['clarity', 'completeness', 'conciseness', 'impact']:
    initial = initial_eval[dim]['score']
    final = best_eval[dim]['score']
    change = final - initial
    print(f'   {dim.title():15} {initial:.2f} → {final:.2f} ({change:+.2f})')

print(f'\n✅ Improvement Complete!')
print(f'   Target: {TARGET_SCORE}+, Achieved: {final_score:.2f}')
if final_score >= TARGET_SCORE:
    print(f'   🎉 TARGET ACHIEVED!')
else:
    print(f'   ⚠️  Target not reached (gap: {TARGET_SCORE - final_score:.2f})')
