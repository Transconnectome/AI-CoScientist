#!/usr/bin/env python3
"""
K-NeuroMind 제안서를 .docx 파일로 생성
한글 프로그램에서 열어 .hwpx로 저장 가능
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_border(table):
    """테이블에 테두리 추가"""
    tbl = table._tbl
    tblPr = tbl.tblPr

    # 테두리 설정
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 1pt = 4
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)

    tblPr.append(tblBorders)

def create_proposal_docx():
    """제안서 .docx 파일 생성"""
    doc = Document()

    # 한글 폰트 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)

    # 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('2026년도 인공지능 분야 신규 R&D 사업 사업보완 제출 양식')
    title_run.font.size = Pt(16)
    title_run.font.bold = True

    doc.add_paragraph()

    # 제출일 및 사업명
    info = doc.add_paragraph()
    info.add_run('제출일: ').bold = True
    info.add_run('2025-10-18\n')
    info.add_run('사업명: ').bold = True
    info.add_run('K-NeuroMind - 한국형 브레인 파운데이션 모델 개발')

    doc.add_page_break()

    # ====================================================================
    # 섹션 1: 사업 내용
    # ====================================================================

    heading = doc.add_heading('1. 사업 내용', level=1)
    heading.runs[0].font.size = Pt(14)

    # 사업 내용 테이블
    table = doc.add_table(rows=11, cols=2)
    table.style = 'Table Grid'
    add_table_border(table)

    # 헤더 행
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '항목'
    hdr_cells[1].text = '내용'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 열 너비 설정
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(5.0)

    # 데이터 입력
    section1_data = [
        ('사업명', 'K-NeuroMind - 한국형 브레인 파운데이션 모델 개발'),
        ('사업 목적',
         '한국은 세계 최고령 사회로 진입하며, 2030년까지 치매 환자가 180만 명으로 2배 증가할 것으로 예상됩니다(보건복지부, 2024). '
         '현재 치매는 증상 발현 후 진단되어 뇌손상이 이미 진행된 상태로 발견되며, 이로 인해 연간 ₩20.7조의 의료비가 소요됩니다.\n\n'
         '본 사업은 이 문제를 해결하기 위해 fMRI, dMRI, EEG 등 다중 모달리티 뇌 데이터를 AI로 통합 분석하여, '
         '"증상 발현 3-5년 전" 질병을 예측하는 "한국형 브레인 파운데이션 모델"을 개발합니다. '
         '이를 통해 조기 개입으로 뇌질환 진행을 지연시키고, 국민 건강 수명을 연장하며, 의료비를 연간 ₩5조 절감하는 것을 목표로 합니다.'),
        ('사업 주요내용',
         '[1단계: 2026-2028년 / 3년] 파운데이션 모델 기반 기술 개발\n\n'
         '1) 다중 모달리티 데이터 통합 인프라 구축\n'
         '• 데이터 규모: 10,000명 (fMRI 8,000명 + dMRI 7,000명 + EEG 5,000명)\n'
         '• 데이터 출처: 기존 KBIS 데이터 80% + 신규 수집 20%\n'
         '• 품질 관리: 자동화 QA 파이프라인\n\n'
         '2) 모달리티별 인코더 개발\n'
         '• fMRI: Vision Transformer (ViT-L/16), 304M 파라미터\n'
         '• dMRI: Graph Attention Network (GAT), 45M 파라미터\n'
         '• EEG: 1D CNN, 18M 파라미터\n'
         '• Validation Accuracy > 80% 달성\n\n'
         '3) 크로스-모달 통합 모델 개발\n'
         '• 아키텍처: Multi-Modal Masked Autoencoder (M³-MAE)\n'
         '• 총 파라미터: 456M\n'
         '• 학습: 64 H100 GPUs × 3개월 (13,824 GPU-hours)\n\n'
         '4) 개념 검증\n'
         '• 알츠하이머 예측 AUC > 0.85 (vs. 기존 SOTA 0.68)\n\n'
         '[2단계: 2029-2030년 / 2년] 모델 고도화 및 실전 배치\n\n'
         '1) 파운데이션 모델 성능 향상\n'
         '• 알츠하이머 예측: AUC 0.85 → 0.92\n'
         '• 다중 질환 확장: 우울증, 조현병, 파킨슨병, 자폐 (각 AUC > 0.80)\n\n'
         '2) 임상 검증\n'
         '• 5개 병원, 1,000명 RCT\n'
         '• 진단 시간 30% 단축, 의료비 20% 절감\n\n'
         '3) K-NeuroMind Open Platform 구축\n'
         '• 공개 모델: 456M params\n'
         '• 공개 데이터: 5,000명\n'
         '• 목표 사용자: 1,000+ 명'),
        ('사업 기간', '2026~2030 (5년)'),
        ('총사업비', '총 101.33억 원 (\'26년 16억원/9개월)'),
        ('\'26년 예산안', '16억 원'),
        ('AI R&D 비중', '총 553.23억 원 (인공지능바이오) 중 16억 원 (2.89%)'),
        ('사업수행 주체', '과학기술정보통신부, 한국연구재단'),
        ('정부 지원 필요성',
         '1. 정부 지원의 필연성\n'
         '• 뇌 데이터 수집: ₩100억\n'
         '• GPU 컴퓨팅: ₩8억\n'
         '• 임상시험: ₩10억\n'
         '→ 총 초기 투자 ₩118억은 민간 단독 감당 불가\n\n'
         '2. 시급성 (Why Now?)\n'
         '• 치매 환자: 84만 → 180만 (2.14배)\n'
         '• 의료비: ₩20.7조/년 → ₩40조/년\n\n'
         '3. 한국의 경쟁우위\n'
         '• KBIS 데이터 5,000명 (세계 3위)\n'
         '• KISTI 슈퍼컴퓨터, NVIDIA 파트너십\n\n'
         '4. 기대 효과\n'
         '• 과학: Nature/Science 10+편\n'
         '• 경제: ₩1.2조 산업 창출\n'
         '• 의료: 연간 ₩5조 절감\n'
         '• 수출: ₩20억/년'),
        ('AI 개발 관련 선행사업', '선행사업 없음')
    ]

    for i, (label, content) in enumerate(section1_data, 1):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = content

    doc.add_page_break()

    # ====================================================================
    # 섹션 2: 인공지능 R&D 특성
    # ====================================================================

    heading = doc.add_heading('2. 인공지능 R&D 특성', level=1)
    heading.runs[0].font.size = Pt(14)

    # 2.1 데이터 수집 및 전처리
    doc.add_heading('2.1 AI 모델 학습을 위한 데이터 수집 및 전처리', level=2)

    # 데이터 수집 방법 테이블
    doc.add_paragraph('▣ 데이터 수집 방법').bold = True

    data_table = doc.add_table(rows=7, cols=2)
    data_table.style = 'Table Grid'
    add_table_border(data_table)

    data_rows = [
        ('데이터 확보 방안', '비율 (%)'),
        ('데이터 신규 구축', '20'),
        ('기존 데이터 확장', '10'),
        ('기 구축된 타기관 데이터 활용', '50'),
        ('기 구축된 자체 데이터 활용', '20'),
        ('기타', '0'),
        ('합계', '100')
    ]

    for i, (item, percent) in enumerate(data_rows):
        row = data_table.rows[i]
        row.cells[0].text = item
        row.cells[1].text = percent
        if i == 0 or i == 6:  # 헤더 및 합계
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph('▣ 데이터 출처·종류·확보방안 등 선택 항목 관련 주요설명').bold = True

    explanation = doc.add_paragraph()
    explanation.add_run('데이터 종류:\n').bold = True
    explanation.add_run(
        '• 뇌영상: fMRI (4D, 91×109×91 voxels), dMRI, T1/T2 MRI\n'
        '• 전기생리: EEG (64 channels, 1000Hz)\n'
        '• 임상: APOE 유전자형, 바이오마커, 인지검사\n\n'
    )
    explanation.add_run('데이터 수집 방법:\n').bold = True
    explanation.add_run(
        '• 5개 병원 (삼성의료원, 서울대병원, 아산병원, 세브란스, 분당서울대병원)\n'
        '• Siemens Prisma 3T 스캐너\n'
        '• IRB 승인 완료\n\n'
    )
    explanation.add_run('신규 구축 비용: ₩20억 (\'26-\'28)\n').bold = True
    explanation.add_run('• 2,000명 × ₩1M = ₩20억\n\n')
    explanation.add_run('기존 데이터 재가공: ₩5억 (\'26-\'27)\n').bold = True

    doc.add_paragraph()

    # 데이터 전처리 테이블
    doc.add_paragraph('▣ 데이터 전처리').bold = True

    preprocess_table = doc.add_table(rows=7, cols=2)
    preprocess_table.style = 'Table Grid'
    add_table_border(preprocess_table)

    preprocess_rows = [
        ('학습·평가 데이터셋 준비', '비율 (%)'),
        ('데이터 기초 전처리', '25'),
        ('데이터 증강 및 변환 적용', '25'),
        ('대량데이터 자동처리 및 최적화', '25'),
        ('실시간 데이터 품질 관리 및 지속적인 갱신', '25'),
        ('기타', '0'),
        ('합계', '100')
    ]

    for i, (item, percent) in enumerate(preprocess_rows):
        row = preprocess_table.rows[i]
        row.cells[0].text = item
        row.cells[1].text = percent
        if i == 0 or i == 6:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph('▣ 데이터전처리 관련 세부 연구내용').bold = True

    preprocess_detail = doc.add_paragraph()
    preprocess_detail.add_run('1단계: 기초 전처리 (25%)\n').bold = True
    preprocess_detail.add_run(
        '• fMRI: FSL FEAT (motion correction, spatial smoothing)\n'
        '• dMRI: FSL eddy (eddy current correction)\n'
        '• EEG: EEGLAB (bandpass filter 0.5-50Hz)\n\n'
    )
    preprocess_detail.add_run('2단계: 데이터 증강 (25%)\n').bold = True
    preprocess_detail.add_run(
        '• 10,000명 → effective 15,000 samples\n'
        '• fMRI: Temporal jittering, Random cropping\n\n'
    )
    preprocess_detail.add_run('3단계: 대량 자동처리 (25%)\n').bold = True
    preprocess_detail.add_run(
        '• KISTI Nurion 500 cores 활용\n'
        '• 1,000명 전처리 → 10일 (vs. 수작업 6개월)\n\n'
    )
    preprocess_detail.add_run('4단계: 품질 관리 (25%)\n').bold = True
    preprocess_detail.add_run(
        '• 실시간 QA (Grafana + Prometheus)\n'
        '• Drift detection (KS test)\n\n'
    )
    preprocess_detail.add_run('전처리 비용: ₩7억 (\'26-\'28)').bold = True

    doc.add_page_break()

    # 2.2 AI 적용/개발 목적
    doc.add_heading('2.2 AI 적용/개발 목적', level=2)

    doc.add_paragraph('▣ 사업 내 AI 연구개발 목적 (택1)').bold = True
    doc.add_paragraph('☑ AI 기술 활용')

    doc.add_paragraph()
    doc.add_paragraph('▣ 선택 항목 관련 주요 내용').bold = True

    purpose = doc.add_paragraph(
        '본 사업은 최첨단 AI 기술(Vision Transformer, Graph Neural Network 등)을 '
        '"바이오·의료 분야"에 적용하여 뇌질환 조기 진단이라는 명확한 사회문제를 해결하는 것이 주 목적입니다.\n\n'
        '세부 내용:\n'
        '1. 적용 분야: 바이오·의료 (뇌질환 진단 및 예측)\n'
        '2. 활용 AI 기술:\n'
        '   • Vision Transformer (fMRI 분석)\n'
        '   • Graph Attention Network (뇌 연결성)\n'
        '   • Multi-Modal Masked Autoencoding\n\n'
        '3. 기대 혁신:\n'
        '   • 기존: 증상 발현 후 진단\n'
        '   • 혁신: 증상 전 3-5년 조기 예측\n'
        '   • 임팩트: 의료비 30% 절감\n\n'
        '4. 산업 파급효과:\n'
        '   • 진단 AI: ₩400억 → ₩1.2조\n'
        '   • BCI 기기: ₩100억 → ₩2.0조'
    )

    doc.add_paragraph()
    doc.add_paragraph('▣ AI 모델 개발 수준 (복수 선택 가능)').bold = True

    model_table = doc.add_table(rows=8, cols=2)
    model_table.style = 'Table Grid'
    add_table_border(model_table)

    model_rows = [
        ('개발 수준', '비율 (%)'),
        ('오픈모델 활용', '50'),
        ('신규모델 개발', '15'),
        ('새로운 알고리즘 개발', '15'),
        ('민간 모델 도입/응용', '10'),
        ('선행사업 모델 활용/개선', '10'),
        ('기타', '0'),
        ('합계', '100')
    ]

    for i, (item, percent) in enumerate(model_rows):
        row = model_table.rows[i]
        row.cells[0].text = item
        row.cells[1].text = percent
        if i == 0 or i == 7:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph('▣ 선택 항목 관련 주요 내용').bold = True

    model_detail = doc.add_paragraph()
    model_detail.add_run('1. 오픈모델 활용 (50%)\n').bold = True
    model_detail.add_run(
        '• PyTorch + Hugging Face Transformers\n'
        '• ViT-L/16 (304M params, ImageNet 사전학습)\n'
        '• Domain adaptation (fine-tuning)\n\n'
    )
    model_detail.add_run('2. 신규모델 개발 (15%)\n').bold = True
    model_detail.add_run(
        '• M³-MAE (Multi-Modal Masked Autoencoder)\n'
        '• 3개 모달리티 동시 처리\n'
        '• 특허 출원 예정 (\'27년)\n\n'
    )
    model_detail.add_run('3. 새로운 알고리즘 개발 (15%)\n').bold = True
    model_detail.add_run(
        '• Loss function 설계\n'
        '• Cross-modal learning\n\n'
    )
    model_detail.add_run('4. 민간 모델 도입 (10%)\n').bold = True
    model_detail.add_run('• OpenAI CLIP, Meta MAE 응용\n\n')
    model_detail.add_run('5. 선행사업 모델 개선 (10%)\n').bold = True
    model_detail.add_run(
        '• BrainIAC 개선\n'
        '• AUC 0.81 → 0.92 목표'
    )

    doc.add_page_break()

    # 2.3 사업 내 AI 활용
    doc.add_heading('2.3 사업 내 AI 활용', level=2)

    doc.add_paragraph('▣ AI 기술 발전 기여도').bold = True

    contribution = doc.add_paragraph()
    contribution.add_run('1. Foundation Model 패러다임의 뇌과학 적용 선도\n').bold = True
    contribution.add_run(
        '• 세계 최초 Multi-Modal Brain Foundation Model\n'
        '• Nature Neuroscience급 논문 3+편 목표\n\n'
    )
    contribution.add_run('2. Self-Supervised Learning의 Small Data 문제 해결\n').bold = True
    contribution.add_run(
        '• 10K 규모에서 효과적 학습 입증\n'
        '• 의료영상 등 Small Data 도메인 적용 가능\n\n'
    )
    contribution.add_run('3. Graph Neural Network의 생물학적 네트워크 적용 확대\n').bold = True
    contribution.add_run(
        '• Brain connectome에 GAT 적용\n'
        '• 단백질 네트워크 등으로 전이 가능\n\n'
    )
    contribution.add_run('4. Cross-Modal Learning 이론 발전\n').bold = True
    contribution.add_run('• fMRI ↔ EEG 상호 예측 최초 검증')

    doc.add_paragraph()
    doc.add_paragraph('▣ AI 보급·확산 기여도').bold = True

    dissemination = doc.add_paragraph()
    dissemination.add_run('1. K-NeuroMind Open Platform (2030년)\n').bold = True
    dissemination.add_run(
        '• 공개 모델: 456M params\n'
        '• 공개 데이터: 5,000명 (50 TB)\n'
        '• KISTI 클라우드 무료 100,000 GPU-hours/년\n'
        '• 목표: 국내 50개 + 해외 100+ 연구자\n\n'
    )
    dissemination.add_run('2. 교육 및 인력 양성\n').bold = True
    dissemination.add_run(
        '• PhD: 10명, 석사: 20명, 인턴: 200명\n'
        '• Coursera/edX 강좌 (목표 10,000명)\n\n'
    )
    dissemination.add_run('3. 산학 협력 생태계\n').bold = True
    dissemination.add_run(
        '• 5개 병원 AI 진단팀\n'
        '• 스타트업 3개 창업\n'
        '• 삼성/LG 공동 연구\n\n'
    )
    dissemination.add_run('4. 국제 협력\n').bold = True
    dissemination.add_run(
        '• NIH BRAIN Initiative 협약\n'
        '• EU Human Brain Project 협력'
    )

    doc.add_paragraph()
    doc.add_paragraph('▣ AI 기술 파급효과').bold = True

    impact = doc.add_paragraph()
    impact.add_run('1. 의료 분야\n').bold = True
    impact.add_run(
        '• 진단 정확도: AUC 0.68 → 0.92 (+35%)\n'
        '• 의료비: 연간 ₩5조 절감\n'
        '• 오진: False negative 40% 감소\n\n'
    )
    impact.add_run('2. 경제 파급효과\n').bold = True
    impact.add_run(
        '• 시장 창출: ₩1.2조 → ₩4조 (CAGR 47%)\n'
        '• 일자리: 직접 80명 + 간접 500명\n'
        '• 수출: ₩20억/년\n\n'
    )
    impact.add_run('3. 사회적 파급효과\n').bold = True
    impact.add_run(
        '• 건강 수명: 치매 free +5년\n'
        '• 삶의 질: QALY +2.6\n'
        '• 간병 부담: 30% 경감\n\n'
    )
    impact.add_run('4. 과학기술 파급효과\n').bold = True
    impact.add_run(
        '• 논문: Nature/Science 2편 + 기타 10편\n'
        '• 특허: 국내 5건, PCT 3건\n'
        '• 국제 위상: AI-뇌과학 Top 3'
    )

    doc.add_page_break()

    # ====================================================================
    # 섹션 3: 주요 기술개발 사항
    # ====================================================================

    heading = doc.add_heading('3. 주요 기술개발 사항', level=1)
    heading.runs[0].font.size = Pt(14)

    doc.add_paragraph('▣ 사업 기간 선택').bold = True
    doc.add_paragraph('☑ 5년 이상 (\'26~\'30)')

    doc.add_paragraph()
    doc.add_paragraph('▣ 3년 이내 기술개발이 어려운 사유').bold = True

    reason = doc.add_paragraph()
    reason.add_run('1. 데이터 수집 기간 (최소 3년)\n').bold = True
    reason.add_run(
        '• Year 1-2: 신규 2,000명 스캔\n'
        '• Year 3: 종단 추적 데이터\n\n'
    )
    reason.add_run('2. 모델 학습 및 검증 기간 (최소 2년)\n').bold = True
    reason.add_run(
        '• Year 3: Pretraining\n'
        '• Year 4: Fine-tuning\n'
        '• Year 4-5: 임상 검증 (RCT 18개월)\n\n'
    )
    reason.add_run('3. 향후 AI 발전 대응 방향\n').bold = True
    reason.add_run(
        '• Modular architecture\n'
        '• Transfer learning\n'
        '• Continuous learning\n'
        '• Explainability 강화'
    )

    doc.add_paragraph()
    doc.add_paragraph('▣ 사업 기간별 주요 기술개발 사항').bold = True

    # 연도별 테이블
    year_table = doc.add_table(rows=6, cols=3)
    year_table.style = 'Table Grid'
    add_table_border(year_table)

    # 열 너비 설정
    year_table.columns[0].width = Inches(0.8)
    year_table.columns[1].width = Inches(3.5)
    year_table.columns[2].width = Inches(2.7)

    year_data = [
        ('연도', '주요 기술개발 사항', '정량적 목표'),
        ('\'26',
         '• 데이터 인프라 구축\n• IRB 승인, BIDS DB\n• QA 파이프라인\n\n• 개별 인코더 개발\n• fMRI: ViT-L/16\n• dMRI: GAT\n• EEG: 1D CNN',
         '• Database v1.0 (3,000명)\n• QA 성공률 > 95%\n\n• fMRI acc > 75%\n• dMRI acc > 70%\n• EEG acc > 70%'),
        ('\'27',
         '• M³-MAE 개발\n• Pretraining (5,000명)\n• Loss 최적화\n\n• 데이터 확장\n• 신규 1,000명\n• Longitudinal 500명',
         '• M³-MAE v1.0 (456M)\n• Masked recon MAE < 0.10\n• Cross-modal acc > 75%\n\n• Total: 6,000명'),
        ('\'28',
         '• 파운데이션 모델 완성\n• Pretraining 완료 (8,000명)\n• 알츠하이머 fine-tuning\n• 국제 벤치마크 테스트',
         '• M³-MAE v2.0\n• Alzheimer AUC > 0.85\n• ADNI: beat SOTA 10%\n• PoC demo 완성'),
        ('\'29',
         '• 다중 질환 모델\n• 5개 질환 fine-tuning\n• BCI 프로토타입\n• RCT 시작 (1,000명)',
         '• 5개 질환 AUC > 0.80\n• BCI < 200ms\n• RCT 중간: 진단 20% 단축'),
        ('\'30',
         '• Open Platform 출시\n• 가중치 공개\n• 데이터셋 공개\n• RCT 완료',
         '• 사용자 1,000+\n• 진단 30% 빠름\n• 비용 20% 감소\n• 3개 스타트업\n• MFDS 인증 신청')
    ]

    for i, (year, dev, goal) in enumerate(year_data):
        row = year_table.rows[i]
        row.cells[0].text = year
        row.cells[1].text = dev
        row.cells[2].text = goal
        if i == 0:
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph('▣ AI 컴퓨팅 자원 구축, 활용 방안').bold = True
    doc.add_paragraph('☑ 클라우드 등을 통한 AI 컴퓨팅 자원 활용')

    doc.add_paragraph()
    doc.add_paragraph('주요 내용:').bold = True

    computing = doc.add_paragraph()
    computing.add_run('1. 컴퓨팅 자원 전략: 하이브리드\n\n').bold = True
    computing.add_run('가. KISTI 슈퍼컴퓨터 (\'26-\'28)\n').bold = True
    computing.add_run(
        '• Nurion (25.7 petaflops, V100)\n'
        '• 무료 500,000 node-hours\n'
        '• 용도: 전처리, 인코더 학습\n\n'
    )
    computing.add_run('나. NVIDIA AI Hub (\'27-\'28)\n').bold = True
    computing.add_run(
        '• 500 H100 × 6개월\n'
        '• M³-MAE pretraining\n'
        '• 비용: ₩41억 → ₩10억 (75% 할인)\n\n'
    )
    computing.add_run('다. Cloud burst (\'28-\'30)\n').bold = True
    computing.add_run(
        '• AWS/GCP GPU instances\n'
        '• Peak demand 대응\n'
        '• 비용: ₩2억/년 × 3년\n\n'
    )
    computing.add_run('라. On-premise GPU (\'28)\n').bold = True
    computing.add_run(
        '• 32 NVIDIA A100\n'
        '• 임상 배포 inference\n'
        '• 비용: ₩8억\n\n'
    )
    computing.add_run('총 컴퓨팅 비용: ₩24억 (5년)').bold = True

    doc.add_page_break()

    # ====================================================================
    # 섹션 4: 예산 규모
    # ====================================================================

    heading = doc.add_heading('4. 예산 규모', level=1)
    heading.runs[0].font.size = Pt(14)

    doc.add_paragraph('▣ AI 관련 예산 규모 (단위: 백만원)').bold = True

    # 예산 테이블
    budget_table = doc.add_table(rows=7, cols=7)
    budget_table.style = 'Table Grid'
    add_table_border(budget_table)

    budget_data = [
        ('항목', '\'26', '\'27', '\'28', '\'29', '\'30', '합계'),
        ('데이터 수집 비용', '500', '600', '400', '200', '100', '1,800'),
        ('데이터 전처리 비용', '200', '250', '200', '50', '50', '750'),
        ('모델 훈련 및 평가\n컴퓨팅 비용', '300', '400', '500', '400', '300', '1,900'),
        ('모델 운영·유지보수 비용', '50', '100', '150', '200', '250', '750'),
        ('기타 AI 관련 비용', '550', '530', '950', '1,515', '1,388', '4,933'),
        ('총계', '1,600', '1,880', '2,200', '2,365', '2,088', '10,133')
    ]

    for i, row_data in enumerate(budget_data):
        row = budget_table.rows[i]
        for j, value in enumerate(row_data):
            cell = row.cells[j]
            cell.text = value
            if i == 0 or i == 6:  # 헤더 및 총계
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif j > 0:  # 수치 열
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    doc.add_paragraph('▣ 세부 설명').bold = True

    budget_detail = doc.add_paragraph()
    budget_detail.add_run('1. 데이터 수집 비용 (₩18억)\n').bold = True
    budget_detail.add_run(
        '• \'26: ₩500M (500명 신규)\n'
        '• \'27: ₩600M (600명 + 종단 추적)\n'
        '• \'28: ₩400M (400명 + 종단 추적)\n'
        '• \'29: ₩200M (임상시험 200명)\n'
        '• \'30: ₩100M (validation)\n\n'
    )
    budget_detail.add_run('2. 데이터 전처리 비용 (₩7.5억)\n').bold = True
    budget_detail.add_run(
        '• \'26: ₩200M (BIDS, QA, 라벨링 500명)\n'
        '• \'27: ₩250M (라벨링 1,000명)\n'
        '• \'28: ₩200M (라벨링 500명, QA)\n'
        '• \'29-\'30: ₩50M (임상 데이터, Platform)\n\n'
    )
    budget_detail.add_run('3. 모델 훈련 및 평가 컴퓨팅 비용 (₩19억)\n').bold = True
    budget_detail.add_run(
        '• \'26: ₩300M (인코더, Cloud)\n'
        '• \'27: ₩400M (Pretraining)\n'
        '• \'28: ₩500M (NVIDIA Hub + On-premise)\n'
        '• \'29: ₩400M (Fine-tuning)\n'
        '• \'30: ₩300M (Continuous learning)\n\n'
    )
    budget_detail.add_run('4. 모델 운영·유지보수 비용 (₩7.5억)\n').bold = True
    budget_detail.add_run(
        '• \'26: ₩50M (인프라)\n'
        '• \'27: ₩100M (MLOps)\n'
        '• \'28: ₩150M (임상 배포)\n'
        '• \'29: ₩200M (5개 병원 운영)\n'
        '• \'30: ₩250M (Platform 운영)\n\n'
    )
    budget_detail.add_run('5. 기타 AI 관련 비용 (₩49.33억)\n').bold = True
    budget_detail.add_run(
        '인건비 + 임상시험 + Platform 개발\n'
        '• \'26: ₩550M (인건비 ₩400M + 장비)\n'
        '• \'27: ₩530M (인건비 + 학회)\n'
        '• \'28: ₩950M (인건비 + 임상 pilot + 플랫폼)\n'
        '• \'29: ₩1,515M (인건비 + RCT + 플랫폼)\n'
        '• \'30: ₩1,388M (인건비 + RCT 완료 + 플랫폼)\n\n'
    )
    budget_detail.add_run('총계 검증: ₩18억 + ₩7.5억 + ₩19억 + ₩7.5억 + ₩49.33억 = ₩101.33억 ✓').bold = True

    # 마지막 페이지
    doc.add_page_break()

    footer = doc.add_paragraph()
    footer.add_run('\n\n문서 작성일: 2025-10-18\n').bold = True
    footer.add_run('작성자: K-NeuroMind 연구팀\n')
    footer.add_run('연락처: [연구책임자 정보]\n')

    # 저장
    output_path = '/Users/jiookcha/Documents/git/AI-CoScientist/output/K-NeuroMind_제안서.docx'
    doc.save(output_path)

    print("=" * 80)
    print("✅ .DOCX 파일 생성 완료!")
    print("=" * 80)
    print()
    print(f"파일 위치: {output_path}")
    print()
    print("다음 단계:")
    print("1. 한글 프로그램 실행")
    print("2. 생성된 .docx 파일 열기")
    print("3. '다른 이름으로 저장' → .hwpx 형식 선택")
    print("4. 저장: K-NeuroMind_제안서_최종.hwpx")
    print()
    print("⚠️  주의: .docx를 열고 최종 검토 후 .hwpx로 저장하세요!")
    print("=" * 80)

if __name__ == "__main__":
    create_proposal_docx()
