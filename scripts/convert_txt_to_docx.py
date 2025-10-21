#!/usr/bin/env python3
"""
Convert improved text file to DOCX format with proper formatting.

This script converts the plain text paper into a properly formatted
Microsoft Word document with:
- Title page
- Section headers
- Proper spacing
- Academic formatting
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def identify_sections(text: str) -> dict:
    """Identify major sections in the paper"""
    sections = {}

    # Common section patterns
    patterns = {
        'title': r'^(.+?)\n',  # First line is usually title
        'abstract': r'(?:ABSTRACT|Abstract|초록)\s*\n(.*?)(?=\n(?:INTRODUCTION|Introduction|서론|1\.))',
        'introduction': r'(?:INTRODUCTION|Introduction|서론|1\.)\s*\n(.*?)(?=\n(?:METHODS|Methods|방법론|2\.))',
        'methods': r'(?:METHODS|Methods|METHODOLOGY|Methodology|방법론|2\.)\s*\n(.*?)(?=\n(?:RESULTS|Results|결과|3\.))',
        'results': r'(?:RESULTS|Results|결과|3\.)\s*\n(.*?)(?=\n(?:DISCUSSION|Discussion|논의|4\.))',
        'discussion': r'(?:DISCUSSION|Discussion|논의|4\.)\s*\n(.*?)(?=\n(?:CONCLUSION|Conclusion|결론|5\.|REFERENCES))',
        'conclusion': r'(?:CONCLUSION|Conclusion|결론|5\.)\s*\n(.*?)(?=\n(?:REFERENCES|References|참고문헌))',
        'references': r'(?:REFERENCES|References|참고문헌)\s*\n(.*)',
    }

    # Extract title (first non-empty line)
    lines = text.split('\n')
    for line in lines:
        if line.strip():
            sections['title'] = line.strip()
            break

    # Extract other sections
    for section_name, pattern in patterns.items():
        if section_name == 'title':
            continue
        match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
        if match:
            sections[section_name] = match.group(1).strip()

    return sections


def create_formatted_document(sections: dict) -> Document:
    """Create a formatted DOCX document from sections"""

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Add title
    if 'title' in sections:
        title = doc.add_heading(sections['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(16)
            run.font.bold = True

    doc.add_paragraph()  # Spacing

    # Add Abstract
    if 'abstract' in sections:
        abstract_heading = doc.add_heading('Abstract', 1)
        for run in abstract_heading.runs:
            run.font.size = Pt(14)
            run.font.bold = True

        abstract_text = sections['abstract']
        # Split into paragraphs
        for para in abstract_text.split('\n\n'):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.size = Pt(12)

        doc.add_paragraph()  # Spacing

    # Add Introduction
    if 'introduction' in sections:
        intro_heading = doc.add_heading('1. Introduction', 1)
        for run in intro_heading.runs:
            run.font.size = Pt(14)
            run.font.bold = True

        intro_text = sections['introduction']
        for para in intro_text.split('\n\n'):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.size = Pt(12)

        doc.add_paragraph()

    # Add Methods
    if 'methods' in sections:
        methods_heading = doc.add_heading('2. Methods', 1)
        for run in methods_heading.runs:
            run.font.size = Pt(14)
            run.font.bold = True

        methods_text = sections['methods']

        # Check for subsections
        subsections = re.split(r'\n(?=\d+\.\d+|\#{2,3}|[A-Z][a-z]+ [A-Z])', methods_text)

        for subsection in subsections:
            if not subsection.strip():
                continue

            # Check if it's a subsection header
            if re.match(r'^\d+\.\d+|^\#{2,3}', subsection):
                lines = subsection.split('\n', 1)
                if len(lines) > 1:
                    doc.add_heading(lines[0].strip(), 2)
                    content = lines[1]
                else:
                    content = subsection
            else:
                content = subsection

            # Add paragraphs
            for para in content.split('\n\n'):
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for run in p.runs:
                        run.font.size = Pt(12)

        doc.add_paragraph()

    # Add Results
    if 'results' in sections:
        results_heading = doc.add_heading('3. Results', 1)
        for run in results_heading.runs:
            run.font.size = Pt(14)
            run.font.bold = True

        results_text = sections['results']
        for para in results_text.split('\n\n'):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.size = Pt(12)

        doc.add_paragraph()

    # Add Discussion
    if 'discussion' in sections:
        discussion_heading = doc.add_heading('4. Discussion', 1)
        for run in discussion_heading.runs:
            run.font.size = Pt(14)
            run.font.bold = True

        discussion_text = sections['discussion']
        for para in discussion_text.split('\n\n'):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.size = Pt(12)

        doc.add_paragraph()

    # Add Conclusion
    if 'conclusion' in sections:
        conclusion_heading = doc.add_heading('5. Conclusion', 1)
        for run in conclusion_heading.runs:
            run.font.size = Pt(14)
            run.font.bold = True

        conclusion_text = sections['conclusion']
        for para in conclusion_text.split('\n\n'):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.size = Pt(12)

        doc.add_paragraph()

    # Add References
    if 'references' in sections:
        refs_heading = doc.add_heading('References', 1)
        for run in refs_heading.runs:
            run.font.size = Pt(14)
            run.font.bold = True

        refs_text = sections['references']

        # Split references (usually numbered or with authors)
        ref_lines = [line.strip() for line in refs_text.split('\n') if line.strip()]

        for ref in ref_lines:
            p = doc.add_paragraph(ref, style='List Number')
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(11)

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    return doc


def convert_txt_to_docx(txt_path: str, docx_path: str = None) -> str:
    """
    Convert text file to DOCX with formatting

    Args:
        txt_path: Path to input text file
        docx_path: Path to output DOCX file (optional)

    Returns:
        Path to created DOCX file
    """
    # Read text file
    with open(txt_path, 'r', encoding='utf-8') as f:
        text = f.read()

    print(f"📄 Reading: {txt_path}")
    print(f"   Length: {len(text):,} characters")

    # Identify sections
    print("\n🔍 Identifying sections...")
    sections = identify_sections(text)
    print(f"   Found sections: {', '.join(sections.keys())}")

    # Create formatted document
    print("\n📝 Creating formatted DOCX...")
    doc = create_formatted_document(sections)

    # Determine output path
    if docx_path is None:
        txt_file = Path(txt_path)
        docx_path = txt_file.with_suffix('.docx')

    # Save document
    doc.save(docx_path)
    print(f"\n✅ Saved: {docx_path}")

    return str(docx_path)


def main():
    """Main entry point"""
    if len(sys.argv) < 2:
        print("Usage: python convert_txt_to_docx.py <input.txt> [output.docx]")
        print("\nExample:")
        print("  python convert_txt_to_docx.py paper_mbbn_improved_final.txt")
        print("  python convert_txt_to_docx.py paper_mbbn_improved_final.txt paper_mbbn_final.docx")
        return 1

    txt_path = sys.argv[1]
    docx_path = sys.argv[2] if len(sys.argv) > 2 else None

    if not Path(txt_path).exists():
        print(f"❌ File not found: {txt_path}")
        return 1

    print("=" * 60)
    print("📄 TXT to DOCX Converter")
    print("=" * 60)

    try:
        output_path = convert_txt_to_docx(txt_path, docx_path)

        print("\n" + "=" * 60)
        print("✅ Conversion Complete")
        print("=" * 60)
        print(f"\nOutput file: {output_path}")

        return 0

    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == '__main__':
    sys.exit(main())
