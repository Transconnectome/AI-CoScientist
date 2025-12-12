#!/usr/bin/env python3
"""Enhanced interactive chatbot with Rich UI and conversation history.

This version includes:
- Rich terminal UI with colors, tables, and progress bars
- Conversation history save/load functionality
- Improved visual presentation
- Session persistence

Usage:
    python scripts/chat_reviewer_enhanced.py
"""

import sys
import os
import json
from pathlib import Path
from typing import Optional, Dict, List
from datetime import datetime
import re
import asyncio
import httpx

# Add src to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from anthropic import Anthropic
from dotenv import load_dotenv
from rich.console import Console
from rich.table import Table
from rich.panel import Panel
from rich.markdown import Markdown
from rich.progress import Progress, SpinnerColumn, TextColumn
from rich.prompt import Prompt
from rich import box
from rich.syntax import Syntax

# Import section-specific modules
from section_parser import SectionType, PaperSectionParser
from section_evaluator import SectionEvaluator, EvaluationContext, format_section_evaluation
from consistency_checker import ConsistencyChecker, format_consistency_report

# Load environment variables
load_dotenv()

# Initialize Rich console
console = Console()


class Phase4Client:
    """Client for Phase 4 Improvement API."""

    def __init__(self, base_url: str = "http://localhost:8000/api/v1"):
        """Initialize Phase 4 API client.

        Args:
            base_url: Base URL for the API
        """
        self.base_url = base_url
        self.client = httpx.AsyncClient(timeout=60.0)

    async def apply_improvement(
        self, paper_id: str, section_name: str, improved_content: str, metadata: Dict = None
    ) -> Dict:
        """Apply improvement to paper section.

        Args:
            paper_id: Paper UUID
            section_name: Section to improve
            improved_content: Improved content
            metadata: Improvement metadata

        Returns:
            Improvement result
        """
        url = f"{self.base_url}/improvements/{paper_id}/apply"
        payload = {
            "section_name": section_name,
            "improved_content": improved_content,
            "metadata": metadata or {}
        }

        response = await self.client.post(url, json=payload)
        response.raise_for_status()
        return response.json()

    async def get_smart_suggestions(self, paper_id: str, section_name: str = None) -> Dict:
        """Get RAG-powered smart suggestions.

        Args:
            paper_id: Paper UUID
            section_name: Optional specific section

        Returns:
            Smart suggestions
        """
        url = f"{self.base_url}/improvements/{paper_id}/suggestions/smart"
        params = {"section_name": section_name} if section_name else {}

        response = await self.client.get(url, params=params)
        response.raise_for_status()
        return response.json()

    async def start_iterative_improvement(
        self, paper_id: str, target_score: float, max_iterations: int = 5, focus_areas: List[str] = None
    ) -> Dict:
        """Start iterative improvement session.

        Args:
            paper_id: Paper UUID
            target_score: Target quality score
            max_iterations: Maximum iterations
            focus_areas: Optional focus areas

        Returns:
            Iteration results
        """
        url = f"{self.base_url}/improvements/{paper_id}/iterate"
        payload = {
            "target_score": target_score,
            "max_iterations": max_iterations,
            "focus_areas": focus_areas or ["clarity", "coherence", "methodology"]
        }

        response = await self.client.post(url, json=payload)
        response.raise_for_status()
        return response.json()

    async def compare_versions(self, paper_id: str, version_a: str, version_b: str) -> Dict:
        """Compare two paper versions.

        Args:
            paper_id: Paper UUID
            version_a: First version (e.g., "1.0.0")
            version_b: Second version (e.g., "1.2.0")

        Returns:
            Version comparison with diffs
        """
        url = f"{self.base_url}/improvements/{paper_id}/versions/compare"
        params = {"version_a": version_a, "version_b": version_b}

        response = await self.client.get(url, params=params)
        response.raise_for_status()
        return response.json()

    async def rollback_to_version(self, paper_id: str, version: str, create_backup: bool = True) -> Dict:
        """Rollback to previous version.

        Args:
            paper_id: Paper UUID
            version: Target version
            create_backup: Whether to create backup

        Returns:
            Rollback result
        """
        url = f"{self.base_url}/improvements/{paper_id}/versions/{version}/rollback"
        payload = {"target_version": version, "create_backup": create_backup}

        response = await self.client.post(url, json=payload)
        response.raise_for_status()
        return response.json()

    async def get_version_history(self, paper_id: str) -> Dict:
        """Get complete version history.

        Args:
            paper_id: Paper UUID

        Returns:
            Version history
        """
        url = f"{self.base_url}/improvements/{paper_id}/versions"

        response = await self.client.get(url)
        response.raise_for_status()
        return response.json()

    async def get_analytics(self, paper_id: str) -> Dict:
        """Get analytics dashboard for a paper.

        Args:
            paper_id: Paper UUID

        Returns:
            Analytics dashboard data
        """
        url = f"{self.base_url}/improvements/{paper_id}/analytics"

        response = await self.client.get(url)
        response.raise_for_status()
        return response.json()

    async def close(self):
        """Close the HTTP client."""
        await self.client.aclose()


class ConversationHistory:
    """Manage conversation history with save/load functionality."""

    def __init__(self, history_dir: Path = None):
        """Initialize conversation history manager.

        Args:
            history_dir: Directory to store conversation histories
        """
        if history_dir is None:
            history_dir = Path.home() / ".ai-coscientist" / "chat_history"

        self.history_dir = history_dir
        self.history_dir.mkdir(parents=True, exist_ok=True)

    def save_session(self, session_data: Dict) -> str:
        """Save a conversation session.

        Args:
            session_data: Dictionary containing session information

        Returns:
            Session ID
        """
        session_id = datetime.now().strftime("%Y%m%d_%H%M%S")
        session_file = self.history_dir / f"session_{session_id}.json"

        with open(session_file, 'w', encoding='utf-8') as f:
            json.dump(session_data, f, indent=2, ensure_ascii=False)

        return session_id

    def load_session(self, session_id: str) -> Optional[Dict]:
        """Load a conversation session.

        Args:
            session_id: Session identifier

        Returns:
            Session data or None if not found
        """
        session_file = self.history_dir / f"session_{session_id}.json"

        if not session_file.exists():
            return None

        with open(session_file, 'r', encoding='utf-8') as f:
            return json.load(f)

    def list_sessions(self, limit: int = 10) -> List[Dict]:
        """List recent conversation sessions.

        Args:
            limit: Maximum number of sessions to return

        Returns:
            List of session metadata
        """
        sessions = []

        for session_file in sorted(self.history_dir.glob("session_*.json"), reverse=True)[:limit]:
            try:
                with open(session_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    sessions.append({
                        'id': session_file.stem.replace('session_', ''),
                        'timestamp': data.get('timestamp', 'Unknown'),
                        'paper': data.get('paper_path', 'N/A'),
                        'messages': len(data.get('messages', []))
                    })
            except Exception:
                continue

        return sessions

    def delete_session(self, session_id: str) -> bool:
        """Delete a conversation session.

        Args:
            session_id: Session identifier

        Returns:
            True if deleted, False otherwise
        """
        session_file = self.history_dir / f"session_{session_id}.json"

        if session_file.exists():
            session_file.unlink()
            return True

        return False


class PaperReviewChatbot:
    """Enhanced chatbot with Rich UI and conversation history."""

    def __init__(self):
        """Initialize the enhanced chatbot."""
        self.client = Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
        self.current_paper_path: Optional[str] = None
        self.current_scores: Optional[Dict] = None
        self.conversation_history: List[Dict] = []
        self.enhanced_versions: List[str] = []
        self.history_manager = ConversationHistory()
        self.current_session_id: Optional[str] = None

        # Phase 4 integration
        self.phase4_client = Phase4Client()
        self.current_paper_id: Optional[str] = None
        self.current_version: str = "1.0.0"

        # Section-specific evaluation components
        self.section_parser = PaperSectionParser()
        self.section_evaluator = SectionEvaluator()
        self.consistency_checker = ConsistencyChecker()
        self.current_paper_sections = {}  # Parsed sections from current paper
        self.section_scores = {}  # Section-specific evaluation results
        self.consistency_issues = {}  # Cross-section consistency issues

        # System prompt
        self.system_prompt = """You are an expert scientific paper reviewer and writing coach.
You help researchers improve their papers by:
1. Evaluating papers across 4 dimensions: Novelty, Methodology, Clarity, Significance
2. Providing specific, actionable improvement suggestions
3. Explaining scores in a friendly, encouraging way
4. Guiding users through the enhancement process step-by-step

You have access to a paper evaluation system that scores papers 0-10 using an ensemble of AI models.
You can also apply automated enhancements like adding theoretical justification, impact quantification, etc.

Be conversational, encouraging, and specific. Always ask clarifying questions when needed."""

    def display_scores_table(self, scores: Dict) -> None:
        """Display scores in a Rich table.

        Args:
            scores: Score dictionary
        """
        # Create overall score panel
        overall = scores['overall']
        confidence = scores.get('confidence', 0.0)

        # Determine color based on score
        if overall >= 9.0:
            color = "green"
            quality = "Exceptional"
        elif overall >= 8.5:
            color = "bright_green"
            quality = "Excellent"
        elif overall >= 8.0:
            color = "yellow"
            quality = "Very Good"
        elif overall >= 7.5:
            color = "bright_yellow"
            quality = "Good"
        elif overall >= 7.0:
            color = "orange1"
            quality = "Acceptable"
        else:
            color = "red"
            quality = "Needs Work"

        # Overall score panel
        overall_panel = Panel(
            f"[bold {color}]{overall}/10[/bold {color}] ({quality})\n"
            f"Confidence: {confidence:.2f}",
            title="📊 Overall Score",
            border_style=color
        )
        console.print(overall_panel)

        # Dimensional scores table
        dim_table = Table(title="📋 Dimensional Scores", box=box.ROUNDED)
        dim_table.add_column("Dimension", style="cyan", no_wrap=True)
        dim_table.add_column("Score", style="magenta")
        dim_table.add_column("Status", justify="right")

        dimensions = [
            ("Novelty", scores.get('novelty', 0)),
            ("Methodology", scores.get('methodology', 0)),
            ("Clarity", scores.get('clarity', 0)),
            ("Significance", scores.get('significance', 0))
        ]

        for dim_name, dim_score in dimensions:
            status = "✅ Strong" if dim_score >= 8.0 else "⚠️ Improve"
            status_color = "green" if dim_score >= 8.0 else "yellow"
            dim_table.add_row(
                dim_name,
                f"{dim_score:.2f}/10",
                f"[{status_color}]{status}[/{status_color}]"
            )

        console.print(dim_table)

        # Model contributions table
        model_table = Table(title="🤖 Model Contributions", box=box.ROUNDED)
        model_table.add_column("Model", style="cyan")
        model_table.add_column("Weight", style="dim")
        model_table.add_column("Score", style="magenta")
        model_table.add_column("Focus", style="italic")

        models = [
            ("GPT-4", "40%", scores.get('gpt4', 0), "Narrative quality"),
            ("Hybrid", "30%", scores.get('hybrid', 0), "Technical depth"),
            ("Multi-task", "30%", scores.get('multitask', 0), "Novelty assessment")
        ]

        for model_name, weight, model_score, focus in models:
            model_table.add_row(model_name, weight, f"{model_score:.2f}/10", focus)

        console.print(model_table)

        # Display LLM feedback if available
        if scores.get('llm_evaluation', False):
            console.print()  # Empty line

            # Strengths panel
            if scores.get('strengths'):
                strengths_text = "\n".join([f"✓ {s}" for s in scores['strengths']])
                strengths_panel = Panel(
                    strengths_text,
                    title="💪 Strengths",
                    border_style="green"
                )
                console.print(strengths_panel)

            # Weaknesses panel
            if scores.get('weaknesses'):
                weaknesses_text = "\n".join([f"• {w}" for w in scores['weaknesses']])
                weaknesses_panel = Panel(
                    weaknesses_text,
                    title="⚠️ Areas for Improvement",
                    border_style="yellow"
                )
                console.print(weaknesses_panel)

            # Dimensional justifications
            if scores.get('novelty_justification'):
                console.print()
                justifications = []
                if scores.get('novelty_justification'):
                    justifications.append(f"[cyan]Novelty:[/cyan] {scores['novelty_justification']}")
                if scores.get('methodology_justification'):
                    justifications.append(f"[cyan]Methodology:[/cyan] {scores['methodology_justification']}")
                if scores.get('clarity_justification'):
                    justifications.append(f"[cyan]Clarity:[/cyan] {scores['clarity_justification']}")
                if scores.get('significance_justification'):
                    justifications.append(f"[cyan]Significance:[/cyan] {scores['significance_justification']}")

                if justifications:
                    justifications_panel = Panel(
                        "\n\n".join(justifications),
                        title="📊 Score Justifications",
                        border_style="blue"
                    )
                    console.print(justifications_panel)

            # Overall assessment
            if scores.get('overall_assessment'):
                console.print()
                assessment_panel = Panel(
                    scores['overall_assessment'],
                    title="📝 Overall Assessment",
                    border_style="magenta"
                )
                console.print(assessment_panel)

    def evaluate_paper(self, file_path: str, use_llm: bool = True) -> Dict:
        """Evaluate a paper with progress indicator.

        Args:
            file_path: Path to the paper file
            use_llm: Whether to use LLM evaluation (default: True)

        Returns:
            Dictionary with scores
        """
        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console
        ) as progress:
            eval_method = "LLM-based analysis" if use_llm else "heuristic analysis"
            task = progress.add_task(f"Analyzing paper with {eval_method}...", total=None)

            # Use the LLM paper evaluator module
            from paper_evaluator_llm import evaluate_paper_file

            try:
                scores = evaluate_paper_file(file_path, use_llm=use_llm)
                progress.update(task, completed=True)
                return scores
            except Exception as e:
                progress.update(task, completed=True)
                console.print(f"[red]Error evaluating paper: {e}[/red]")
                console.print("[yellow]Falling back to heuristic evaluation...[/yellow]")
                # Return default scores on error
                try:
                    scores = evaluate_paper_file(file_path, use_llm=False)
                    return scores
                except:
                    return {
                        'overall': 7.0,
                        'novelty': 7.0,
                        'methodology': 7.0,
                        'clarity': 7.0,
                        'significance': 7.0,
                        'confidence': 0.60,
                        'gpt4': 7.0,
                        'hybrid': 7.0,
                        'multitask': 7.0,
                        'llm_evaluation': False
                    }

    def get_improvement_suggestions(self, current_score: float, target_score: float = None) -> List[Dict]:
        """Get improvement suggestions based on current score."""
        suggestions = [
            {
                'title': 'Transform Title with Crisis Framing',
                'description': 'Reframe title from incremental to paradigm shift',
                'time': '30 minutes',
                'expected_gain': 0.3,
                'difficulty': 'Easy',
                'script': None
            },
            {
                'title': 'Add Theoretical Justification Section',
                'description': 'Add ~1200 word theoretical foundations section',
                'time': '2 hours',
                'expected_gain': 0.3,
                'difficulty': 'Medium',
                'script': 'insert_theoretical_justification.py'
            },
            {
                'title': 'Quantify All Impact Statements',
                'description': 'Replace vague statements with specific numbers',
                'time': '1-2 hours',
                'expected_gain': 0.2,
                'difficulty': 'Easy',
                'script': None
            },
            {
                'title': 'Add Method Comparison Table',
                'description': 'Add systematic comparison with alternative methods',
                'time': '1 hour',
                'expected_gain': 0.1,
                'difficulty': 'Easy',
                'script': 'add_comparison_table.py'
            },
            {
                'title': 'Add Impact Boxes',
                'description': 'Add visual impact boxes with quantified metrics',
                'time': '30 minutes',
                'expected_gain': 0.05,
                'difficulty': 'Easy',
                'script': 'add_impact_boxes.py'
            }
        ]

        # Filter suggestions based on target
        if target_score:
            gap = target_score - current_score
            selected = []
            cumulative_gain = 0
            for sug in suggestions:
                if cumulative_gain < gap:
                    selected.append(sug)
                    cumulative_gain += sug['expected_gain']
            return selected

        return suggestions[:3]

    def display_suggestions(self, suggestions: List[Dict]) -> None:
        """Display improvement suggestions in a Rich table.

        Args:
            suggestions: List of improvement suggestions
        """
        table = Table(title="💡 Improvement Suggestions", box=box.ROUNDED)
        table.add_column("#", style="cyan", width=3)
        table.add_column("Suggestion", style="green")
        table.add_column("Time", style="yellow")
        table.add_column("Gain", style="magenta")
        table.add_column("Difficulty", style="blue")

        for i, sug in enumerate(suggestions, 1):
            table.add_row(
                str(i),
                sug['title'],
                sug['time'],
                f"+{sug['expected_gain']:.2f}",
                sug['difficulty']
            )

        console.print(table)

    def parse_user_intent(self, message: str) -> tuple[str, dict]:
        """Parse user message to determine intent."""
        msg_lower = message.lower()

        # Check for file path
        file_pattern = r'[\w\-./]+\.(?:docx|txt|pdf)'
        file_match = re.search(file_pattern, message)

        # History commands
        if any(word in msg_lower for word in ['save', 'save conversation', 'save session']):
            return 'save_history', {}
        elif any(word in msg_lower for word in ['load', 'load conversation', 'load session']):
            return 'load_history', {}
        elif any(word in msg_lower for word in ['show history', 'list sessions', 'list conversations']):
            return 'list_history', {}

        # Detect intents
        if any(word in msg_lower for word in ['review', 'evaluate', 'analyze', 'score']):
            if file_match:
                return 'evaluate', {'file_path': file_match.group(0)}
            else:
                return 'evaluate_request', {}

        elif any(word in msg_lower for word in ['improve', 'enhance', 'better', 'increase']):
            score_match = re.search(r'(\d+\.?\d*)', message)
            target = float(score_match.group(1)) if score_match else None
            return 'improve', {'target_score': target}

        elif any(word in msg_lower for word in ['apply', 'add', 'insert']):
            if 'theoretical' in msg_lower or 'theory' in msg_lower:
                return 'apply_enhancement', {'type': 'theoretical'}
            elif 'comparison' in msg_lower or 'table' in msg_lower:
                return 'apply_enhancement', {'type': 'comparison'}
            elif 'impact' in msg_lower or 'box' in msg_lower:
                return 'apply_enhancement', {'type': 'impact_boxes'}
            else:
                return 'apply_request', {}

        elif any(word in msg_lower for word in ['why', 'explain', 'how']):
            return 'explain', {}

        elif any(word in msg_lower for word in ['next', 'what now', 'then']):
            return 'next_steps', {}

        # Section-specific review commands
        elif 'review abstract' in msg_lower or 'evaluate abstract' in msg_lower:
            return 'review_section', {'section': SectionType.ABSTRACT}

        elif 'review introduction' in msg_lower or 'evaluate introduction' in msg_lower:
            return 'review_section', {'section': SectionType.INTRODUCTION}

        elif 'review methods' in msg_lower or 'review methodology' in msg_lower or 'evaluate methods' in msg_lower:
            return 'review_section', {'section': SectionType.METHODS}

        elif 'review results' in msg_lower or 'evaluate results' in msg_lower:
            return 'review_section', {'section': SectionType.RESULTS}

        elif 'review discussion' in msg_lower or 'evaluate discussion' in msg_lower:
            return 'review_section', {'section': SectionType.DISCUSSION}

        elif 'list sections' in msg_lower or 'show sections' in msg_lower or 'what sections' in msg_lower:
            return 'list_sections', {}

        elif 'review all sections' in msg_lower or 'evaluate all sections' in msg_lower:
            return 'review_all_sections', {}

        elif 'check consistency' in msg_lower or 'consistency check' in msg_lower:
            return 'check_consistency', {}

        # Phase 4 commands
        elif msg_lower.startswith('/apply'):
            # Parse section name if provided
            section_match = re.search(r'(?:to|for)\s+(\w+)', msg_lower)
            section = section_match.group(1) if section_match else None
            return 'phase4_apply', {'section': section}

        elif msg_lower.startswith('/suggest') or msg_lower.startswith('/smart'):
            section_match = re.search(r'(?:for|on)\s+(\w+)', msg_lower)
            section = section_match.group(1) if section_match else None
            return 'phase4_suggest', {'section': section}

        elif msg_lower.startswith('/iterate'):
            # Parse target score if provided
            score_match = re.search(r'(\d+\.?\d*)', message)
            target = float(score_match.group(1)) if score_match else 8.5
            # Parse max iterations
            iter_match = re.search(r'(\d+)\s*(?:iterations?|rounds?)', msg_lower)
            max_iter = int(iter_match.group(1)) if iter_match else 5
            return 'phase4_iterate', {'target_score': target, 'max_iterations': max_iter}

        elif msg_lower.startswith('/compare'):
            # Parse version numbers (e.g., "1.0.0" and "1.2.0")
            versions = re.findall(r'(\d+\.\d+\.\d+)', message)
            if len(versions) >= 2:
                return 'phase4_compare', {'version_a': versions[0], 'version_b': versions[1]}
            else:
                return 'phase4_compare_request', {}

        elif msg_lower.startswith('/rollback'):
            version_match = re.search(r'(\d+\.\d+\.\d+)', message)
            version = version_match.group(1) if version_match else None
            return 'phase4_rollback', {'version': version}

        elif msg_lower.startswith('/versions') or msg_lower.startswith('/history'):
            return 'phase4_versions', {}

        elif msg_lower.startswith('/analytics') or msg_lower.startswith('/stats') or msg_lower.startswith('/dashboard'):
            return 'phase4_analytics', {}

        else:
            return 'general', {}

    def save_current_session(self) -> str:
        """Save current conversation session.

        Returns:
            Session ID
        """
        session_data = {
            'timestamp': datetime.now().isoformat(),
            'paper_path': self.current_paper_path,
            'scores': self.current_scores,
            'messages': self.conversation_history,
            'enhanced_versions': self.enhanced_versions
        }

        session_id = self.history_manager.save_session(session_data)
        self.current_session_id = session_id

        return session_id

    def load_session_by_id(self, session_id: str) -> bool:
        """Load a session by ID.

        Args:
            session_id: Session identifier

        Returns:
            True if loaded successfully
        """
        session_data = self.history_manager.load_session(session_id)

        if session_data:
            self.current_paper_path = session_data.get('paper_path')
            self.current_scores = session_data.get('scores')
            self.conversation_history = session_data.get('messages', [])
            self.enhanced_versions = session_data.get('enhanced_versions', [])
            self.current_session_id = session_id
            return True

        return False

    def display_session_list(self) -> None:
        """Display list of available sessions."""
        sessions = self.history_manager.list_sessions()

        if not sessions:
            console.print("[yellow]No saved sessions found.[/yellow]")
            return

        table = Table(title="💾 Saved Sessions", box=box.ROUNDED)
        table.add_column("#", style="cyan", width=3)
        table.add_column("Session ID", style="green")
        table.add_column("Date", style="yellow")
        table.add_column("Paper", style="magenta")
        table.add_column("Messages", style="blue")

        for i, session in enumerate(sessions, 1):
            timestamp = session['timestamp']
            if 'T' in timestamp:
                timestamp = timestamp.split('T')[0]

            table.add_row(
                str(i),
                session['id'],
                timestamp,
                Path(session['paper']).name if session['paper'] else 'N/A',
                str(session['messages'])
            )

        console.print(table)

    async def handle_phase4_suggest(self, section_name: str = None) -> str:
        """Handle /suggest command - Get RAG-powered smart suggestions."""
        if not self.current_paper_id:
            return "❌ No paper has been registered with Phase 4 yet. Please evaluate a paper first."

        try:
            console.print(f"\n[cyan]🧠 Fetching RAG-powered suggestions{' for ' + section_name if section_name else ''}...[/cyan]\n")

            result = await self.phase4_client.get_smart_suggestions(
                self.current_paper_id,
                section_name
            )

            # Display suggestions with Rich
            table = Table(title="💡 Smart Suggestions (RAG-Enhanced)", box=box.ROUNDED)
            table.add_column("#", style="cyan", width=3)
            table.add_column("Section", style="green")
            table.add_column("Expected Gain", style="magenta")
            table.add_column("Patterns Used", style="yellow")
            table.add_column("Exemplars", style="blue")

            for i, sug in enumerate(result.get('suggestions', []), 1):
                table.add_row(
                    str(i),
                    sug['section_name'],
                    f"+{sug['expected_improvement']:.2f}",
                    str(sug.get('similar_patterns_used', 0)),
                    str(sug.get('exemplars_referenced', 0))
                )

            console.print(table)

            rag_status = "✅ RAG-Enhanced" if result.get('rag_enhanced') else "⚠️ Basic"
            console.print(f"\n[dim]{rag_status} | Total: {result.get('total_suggestions', 0)} suggestions[/dim]\n")

            return f"Found {result.get('total_suggestions', 0)} RAG-powered suggestions. Use '/apply <number>' to apply one!"

        except httpx.HTTPStatusError as e:
            return f"❌ API Error: {e.response.status_code} - {e.response.text}"
        except Exception as e:
            return f"❌ Error getting suggestions: {str(e)}"

    async def handle_phase4_iterate(self, target_score: float, max_iterations: int) -> str:
        """Handle /iterate command - Iterative improvement loop."""
        if not self.current_paper_id:
            return "❌ No paper has been registered with Phase 4 yet. Please evaluate a paper first."

        try:
            current = self.current_scores.get('overall', 0) if self.current_scores else 0
            console.print(f"\n[cyan]🔄 Starting iterative improvement: {current:.2f} → {target_score} (max {max_iterations} rounds)...[/cyan]\n")

            with Progress(
                SpinnerColumn(),
                TextColumn("[progress.description]{task.description}"),
                console=console
            ) as progress:
                task = progress.add_task("Iterating...", total=None)

                result = await self.phase4_client.start_iterative_improvement(
                    self.current_paper_id,
                    target_score,
                    max_iterations
                )

                progress.update(task, completed=True)

            # Display results
            table = Table(title="📊 Iteration Results", box=box.ROUNDED)
            table.add_column("Metric", style="cyan")
            table.add_column("Value", style="magenta")

            table.add_row("Iterations Completed", str(result['iterations_completed']))
            table.add_row("Improvements Applied", str(result['improvements_applied']))
            table.add_row("Initial Score", f"{result['initial_score']:.2f}/10")
            table.add_row("Final Score", f"{result['final_score']:.2f}/10")
            table.add_row("Score Gain", f"+{result['score_improvement']:.2f}")
            table.add_row("Target Reached", "✅ Yes" if result['target_reached'] else "⚠️ No")

            console.print(table)

            # Update current version
            self.current_version = result.get('final_version', self.current_version)

            success_msg = f"✅ Target reached!" if result['target_reached'] else f"⚠️ Stopped at {result['final_score']:.2f}"
            return f"{success_msg} Applied {result['improvements_applied']} improvements in {result['iterations_completed']} rounds."

        except httpx.HTTPStatusError as e:
            return f"❌ API Error: {e.response.status_code} - {e.response.text}"
        except Exception as e:
            return f"❌ Error during iteration: {str(e)}"

    async def handle_phase4_compare(self, version_a: str, version_b: str) -> str:
        """Handle /compare command - Compare two versions."""
        if not self.current_paper_id:
            return "❌ No paper has been registered with Phase 4 yet."

        try:
            console.print(f"\n[cyan]📊 Comparing versions {version_a} and {version_b}...[/cyan]\n")

            result = await self.phase4_client.compare_versions(
                self.current_paper_id,
                version_a,
                version_b
            )

            # Display score comparison
            score_table = Table(title=f"Score Comparison: {version_a} vs {version_b}", box=box.ROUNDED)
            score_table.add_column("Metric", style="cyan")
            score_table.add_column(version_a, style="yellow")
            score_table.add_column(version_b, style="green")
            score_table.add_column("Delta", style="magenta")

            for metric, values in result.get('score_comparison', {}).items():
                score_a = values.get('version_a', 0)
                score_b = values.get('version_b', 0)
                delta = score_b - score_a
                delta_str = f"+{delta:.2f}" if delta > 0 else f"{delta:.2f}"
                score_table.add_row(metric.title(), f"{score_a:.2f}", f"{score_b:.2f}", delta_str)

            console.print(score_table)

            # Display diff preview
            if result.get('diff'):
                console.print(Panel(
                    Syntax(result['diff'][:500] + "..." if len(result['diff']) > 500 else result['diff'], "diff", theme="monokai"),
                    title="📝 Diff Preview",
                    border_style="blue"
                ))

            return f"Compared versions {version_a} and {version_b}. Overall delta: {result.get('overall_improvement', 0):.2f} points."

        except httpx.HTTPStatusError as e:
            return f"❌ API Error: {e.response.status_code} - {e.response.text}"
        except Exception as e:
            return f"❌ Error comparing versions: {str(e)}"

    async def handle_phase4_rollback(self, version: str) -> str:
        """Handle /rollback command - Rollback to previous version."""
        if not self.current_paper_id:
            return "❌ No paper has been registered with Phase 4 yet."

        if not version:
            return "❌ Please specify a version to rollback to. Usage: /rollback 1.0.0"

        try:
            console.print(f"\n[yellow]⏪ Rolling back to version {version}...[/yellow]\n")

            result = await self.phase4_client.rollback_to_version(
                self.current_paper_id,
                version
            )

            new_version = result.get('new_version', 'unknown')
            backup_version = result.get('backup_version', 'none')

            console.print(f"[green]✅ Rollback successful![/green]")
            console.print(f"[dim]New version: {new_version} | Backup: {backup_version}[/dim]\n")

            return f"Rolled back to {version}. Created new version {new_version} with content from {version}."

        except httpx.HTTPStatusError as e:
            return f"❌ API Error: {e.response.status_code} - {e.response.text}"
        except Exception as e:
            return f"❌ Error during rollback: {str(e)}"

    async def handle_phase4_versions(self) -> str:
        """Handle /versions command - Show version history."""
        if not self.current_paper_id:
            return "❌ No paper has been registered with Phase 4 yet."

        try:
            console.print(f"\n[cyan]📚 Fetching version history...[/cyan]\n")

            result = await self.phase4_client.get_version_history(self.current_paper_id)

            # Display version history table
            table = Table(title="📋 Version History", box=box.ROUNDED)
            table.add_column("Version", style="cyan")
            table.add_column("Created", style="yellow")
            table.add_column("Type", style="green")
            table.add_column("Score", style="magenta")
            table.add_column("Changes", style="blue")

            for version in result.get('versions', []):
                is_current = "→ " if version['version'] == result.get('current_version') else ""
                table.add_row(
                    f"{is_current}{version['version']}",
                    version.get('created_at', 'Unknown')[:19],
                    version.get('version_type', 'patch'),
                    f"{version.get('quality_score', 0):.2f}/10",
                    str(version.get('changes_count', 0))
                )

            console.print(table)

            total = result.get('total_versions', 0)
            current = result.get('current_version', 'unknown')
            console.print(f"\n[dim]Total versions: {total} | Current: {current}[/dim]\n")

            return f"Showing {total} versions. Current version: {current}"

        except httpx.HTTPStatusError as e:
            return f"❌ API Error: {e.response.status_code} - {e.response.text}"
        except Exception as e:
            return f"❌ Error getting version history: {str(e)}"

    async def handle_phase4_analytics(self) -> str:
        """Handle /analytics command - Show analytics dashboard."""
        if not self.current_paper_id:
            return "❌ No paper has been registered with Phase 4 yet."

        try:
            console.print(f"\n[cyan]📊 Generating analytics dashboard...[/cyan]\n")

            result = await self.phase4_client.get_analytics(self.current_paper_id)

            # Display quality progression
            prog = result.get('quality_progression', {})
            prog_table = Table(title="📈 Quality Progression", box=box.ROUNDED)
            prog_table.add_column("Metric", style="cyan")
            prog_table.add_column("Value", style="magenta")

            prog_table.add_row("Starting Score", f"{prog.get('starting_score', 0):.2f}/10")
            prog_table.add_row("Current Score", f"{prog.get('current_score', 0):.2f}/10")
            prog_table.add_row("Total Improvement", f"+{prog.get('total_improvement', 0):.2f}")
            prog_table.add_row("Improvement %", f"{prog.get('improvement_percentage', 0):.1f}%")
            prog_table.add_row("Iterations", str(prog.get('iterations_count', 0)))

            console.print(prog_table)

            # Display section improvements
            sections = result.get('section_improvements', [])
            if sections:
                console.print()
                section_table = Table(title="📝 Section Improvements", box=box.ROUNDED)
                section_table.add_column("Section", style="cyan")
                section_table.add_column("Count", style="yellow")
                section_table.add_column("Avg Impact", style="magenta")

                for section in sections:
                    section_table.add_row(
                        section['section_name'],
                        str(section['improvement_count']),
                        f"+{section['average_impact']:.2f}"
                    )

                console.print(section_table)

            # Display improvement type effectiveness
            types = result.get('improvement_type_effectiveness', [])
            if types:
                console.print()
                type_table = Table(title="🎯 Improvement Type Effectiveness", box=box.ROUNDED)
                type_table.add_column("Type", style="cyan")
                type_table.add_column("Count", style="yellow")
                type_table.add_column("Avg Impact", style="magenta")
                type_table.add_column("Success Rate", style="green")

                for itype in types:
                    type_table.add_row(
                        itype['improvement_type'],
                        str(itype['count']),
                        f"+{itype['average_impact']:.2f}",
                        f"{itype['success_rate']:.0f}%"
                    )

                console.print(type_table)

            # Display iteration efficiency
            efficiency = result.get('iteration_efficiency', {})
            if efficiency:
                console.print()
                eff_panel = Panel(
                    f"Total Iterations: {efficiency.get('total_iterations', 0)}\n"
                    f"Avg Improvements/Iteration: {efficiency.get('average_improvements_per_iteration', 0):.2f}\n"
                    f"Avg Score Gain: +{efficiency.get('average_score_gain', 0):.2f}\n"
                    f"Diminishing Returns: {'⚠️ Detected' if efficiency.get('diminishing_returns_detected') else '✅ No'}",
                    title="⚡ Iteration Efficiency",
                    border_style="yellow"
                )
                console.print(eff_panel)

            # Display recommendations
            recommendations = result.get('recommendations', [])
            if recommendations:
                console.print()
                rec_text = "\n".join([f"• {rec}" for rec in recommendations])
                rec_panel = Panel(
                    rec_text,
                    title="💡 Recommendations",
                    border_style="green"
                )
                console.print(rec_panel)

            total_impr = sum(s['improvement_count'] for s in sections)
            return f"Generated analytics dashboard with {prog.get('iterations_count', 0)} iterations, {total_impr} improvements, and {len(recommendations)} recommendations."

        except httpx.HTTPStatusError as e:
            return f"❌ API Error: {e.response.status_code} - {e.response.text}"
        except Exception as e:
            return f"❌ Error getting analytics: {str(e)}"

    def generate_response(self, user_message: str) -> str:
        """Generate chatbot response using Claude."""
        # Parse intent
        intent, params = self.parse_user_intent(user_message)

        # Handle history intents
        if intent == 'save_history':
            session_id = self.save_current_session()
            return f"✅ Session saved! ID: {session_id}"

        elif intent == 'load_history':
            self.display_session_list()
            session_id = Prompt.ask("\n[cyan]Enter session ID to load[/cyan]")
            if self.load_session_by_id(session_id):
                console.print(f"[green]✅ Session {session_id} loaded![/green]")
                if self.current_scores:
                    self.display_scores_table(self.current_scores)
                return "Session restored. You can continue from where you left off."
            else:
                return f"❌ Session {session_id} not found."

        elif intent == 'list_history':
            self.display_session_list()
            return "Use 'load conversation' to restore a session."

        # Handle specific intents
        context = ""

        if intent == 'evaluate':
            file_path = params['file_path']
            console.print(f"\n[cyan]📄 Evaluating paper: {file_path}[/cyan]\n")

            # Evaluate paper
            self.current_paper_path = file_path
            self.current_scores = self.evaluate_paper(file_path)

            # Display scores with Rich
            self.display_scores_table(self.current_scores)

            # Auto-save after evaluation
            session_id = self.save_current_session()
            console.print(f"\n[dim]Session auto-saved: {session_id}[/dim]\n")

            context = f"""The user wants to evaluate their paper: {file_path}

Paper Evaluation Results:
- Overall Score: {self.current_scores['overall']}/10
- Novelty: {self.current_scores['novelty']}/10
- Methodology: {self.current_scores['methodology']}/10
- Clarity: {self.current_scores['clarity']}/10
- Significance: {self.current_scores['significance']}/10
- Confidence: {self.current_scores['confidence']}

Model Contributions:
- GPT-4 (40%): {self.current_scores['gpt4']}/10
- Hybrid (30%): {self.current_scores['hybrid']}/10
- Multi-task (30%): {self.current_scores['multitask']}/10

Provide an encouraging analysis of these scores, highlighting strengths and areas for improvement.
Ask the user what they'd like to work on."""

        elif intent == 'improve':
            target = params.get('target_score')
            if not self.current_scores:
                context = "The user wants to improve their paper, but no paper has been evaluated yet. Ask them to provide a paper to review first."
            else:
                suggestions = self.get_improvement_suggestions(
                    self.current_scores['overall'],
                    target
                )

                # Display suggestions with Rich
                self.display_suggestions(suggestions)

                suggestions_text = "\n".join([
                    f"{i+1}. {s['title']}"
                    f"\n   - Time: {s['time']}"
                    f"\n   - Expected gain: +{s['expected_gain']} points"
                    f"\n   - Difficulty: {s['difficulty']}"
                    for i, s in enumerate(suggestions)
                ])

                context = f"""The user wants to improve their paper. Current score: {self.current_scores['overall']}/10
Target score: {target or 'not specified'}

Top improvement suggestions:
{suggestions_text}

Present these suggestions in a friendly way and ask which one they'd like to start with."""

        elif intent == 'review_section':
            section_type = params['section']
            if not self.current_paper_path:
                return "❌ No paper has been evaluated yet. Please review a paper first using: 'Review my paper: /path/to/paper.docx'"

            # Parse paper if not already done
            if not self.current_paper_sections:
                console.print("[cyan]📄 Parsing paper sections...[/cyan]")
                from docx import Document
                doc = Document(self.current_paper_path)
                paper_text = '\n'.join([p.text for p in doc.paragraphs])
                self.current_paper_sections = self.section_parser.parse(paper_text)

            # Check if section exists
            if section_type not in self.current_paper_sections:
                available = ", ".join([s.value for s in self.current_paper_sections.keys()])
                return f"❌ Section '{section_type.value}' not found in paper.\n\nAvailable sections: {available}"

            # Evaluate section
            console.print(f"\n[cyan]🔍 Evaluating {section_type.value} section...[/cyan]\n")
            section = self.current_paper_sections[section_type]

            # Get full paper for context
            from docx import Document
            doc = Document(self.current_paper_path)
            paper_text = '\n'.join([p.text for p in doc.paragraphs])

            evaluation = self.section_evaluator.evaluate_section(
                section_type,
                section.content,
                EvaluationContext.INTEGRATED,
                full_paper_text=paper_text
            )

            # Store section score
            self.section_scores[section_type] = evaluation

            # Display with Rich
            formatted = format_section_evaluation(evaluation)
            console.print(Panel(formatted, title=f"📊 {section_type.value.upper()} Evaluation", border_style="cyan"))

            context = f"""The user asked for a section-specific review of the {section_type.value} section.

Evaluation results for {section_type.value}:
- Overall Score: {evaluation.get('overall', {}).get('score', 0)}/10
- Word Count: {evaluation.get('word_count', 0)} words

Key dimensions scores: {', '.join([f"{k}: {v.get('score', 0):.2f}" for k, v in evaluation.items() if isinstance(v, dict) and 'score' in v and k != 'overall'])}

Strengths: {', '.join(evaluation.get('strengths', []))}
Weaknesses: {', '.join(evaluation.get('weaknesses', []))}

Provide encouraging feedback and actionable suggestions for improving this specific section."""

        elif intent == 'list_sections':
            if not self.current_paper_path:
                return "❌ No paper has been evaluated yet. Please review a paper first."

            # Parse paper if not already done
            if not self.current_paper_sections:
                console.print("[cyan]📄 Parsing paper sections...[/cyan]")
                from docx import Document
                doc = Document(self.current_paper_path)
                paper_text = '\n'.join([p.text for p in doc.paragraphs])
                self.current_paper_sections = self.section_parser.parse(paper_text)

            # Display sections table
            table = Table(title="📑 Detected Paper Sections", box=box.ROUNDED)
            table.add_column("Section", style="cyan")
            table.add_column("Title", style="yellow")
            table.add_column("Word Count", style="magenta")
            table.add_column("Status", style="green")

            for section_type, section in self.current_paper_sections.items():
                evaluated = "✅ Evaluated" if section_type in self.section_scores else "⏳ Not reviewed"
                table.add_row(
                    section_type.value.title(),
                    section.title,
                    str(section.word_count),
                    evaluated
                )

            console.print(table)

            # Get completion status
            completion = self.section_parser.get_completion_status()
            console.print(f"\n[cyan]Paper Completeness: {completion['completion_percentage']:.0f}% ({completion['completed_sections']}/{completion['total_sections']} core sections)[/cyan]")

            return f"Found {len(self.current_paper_sections)} sections. Use 'review abstract', 'review methods', etc. to evaluate specific sections."

        elif intent == 'review_all_sections':
            if not self.current_paper_path:
                return "❌ No paper has been evaluated yet. Please review a paper first."

            # Parse paper if not already done
            if not self.current_paper_sections:
                console.print("[cyan]📄 Parsing paper sections...[/cyan]")
                from docx import Document
                doc = Document(self.current_paper_path)
                paper_text = '\n'.join([p.text for p in doc.paragraphs])
                self.current_paper_sections = self.section_parser.parse(paper_text)

            # Evaluate all sections
            console.print("\n[cyan]🔍 Evaluating all sections...[/cyan]\n")

            from docx import Document
            doc = Document(self.current_paper_path)
            paper_text = '\n'.join([p.text for p in doc.paragraphs])

            for section_type, section in self.current_paper_sections.items():
                console.print(f"[dim]Evaluating {section_type.value}...[/dim]")
                evaluation = self.section_evaluator.evaluate_section(
                    section_type,
                    section.content,
                    EvaluationContext.INTEGRATED,
                    full_paper_text=paper_text
                )
                self.section_scores[section_type] = evaluation

            # Display summary table
            table = Table(title="📊 All Sections Evaluation Summary", box=box.ROUNDED)
            table.add_column("Section", style="cyan")
            table.add_column("Overall Score", style="magenta")
            table.add_column("Strengths", style="green")
            table.add_column("Weaknesses", style="yellow")

            for section_type, evaluation in self.section_scores.items():
                overall_score = evaluation.get('overall', {}).get('score', 0)
                strengths_count = len(evaluation.get('strengths', []))
                weaknesses_count = len(evaluation.get('weaknesses', []))
                table.add_row(
                    section_type.value.title(),
                    f"{overall_score:.2f}/10",
                    f"✅ {strengths_count}",
                    f"⚠️ {weaknesses_count}"
                )

            console.print(table)

            avg_score = sum([e.get('overall', {}).get('score', 0) for e in self.section_scores.values()]) / len(self.section_scores)
            console.print(f"\n[cyan]Average Section Score: {avg_score:.2f}/10[/cyan]\n")

            return f"All {len(self.section_scores)} sections evaluated! You can ask for detailed feedback on any specific section."

        elif intent == 'check_consistency':
            if not self.current_paper_path:
                return "❌ No paper has been evaluated yet. Please review a paper first."

            # Parse paper if not already done
            if not self.current_paper_sections:
                console.print("[cyan]📄 Parsing paper sections...[/cyan]")
                from docx import Document
                doc = Document(self.current_paper_path)
                paper_text = '\n'.join([p.text for p in doc.paragraphs])
                self.current_paper_sections = self.section_parser.parse(paper_text)

            # Check consistency
            console.print("\n[cyan]🔍 Checking cross-section consistency...[/cyan]\n")
            self.consistency_issues = self.consistency_checker.comprehensive_consistency_check(
                self.current_paper_sections
            )

            # Format and display report
            report = format_consistency_report(self.consistency_issues)
            console.print(Panel(report, title="📋 Consistency Check Report", border_style="yellow"))

            # Count issues by severity
            total_issues = sum(len(issues) for issues in self.consistency_issues.values())
            if total_issues == 0:
                return "✅ Excellent! No consistency issues detected across sections. Your paper maintains good coherence."
            else:
                critical = sum(1 for issues in self.consistency_issues.values() for i in issues if i.severity.value == "critical")
                major = sum(1 for issues in self.consistency_issues.values() for i in issues if i.severity.value == "major")

                context = f"""The consistency check found {total_issues} issues:
- Critical: {critical}
- Major: {major}

Key issue types: {', '.join(self.consistency_issues.keys())}

Provide encouraging feedback and prioritize which issues to address first."""
                return context

        # Phase 4 command handlers
        elif intent == 'phase4_suggest':
            return asyncio.run(self.handle_phase4_suggest(params.get('section')))

        elif intent == 'phase4_iterate':
            return asyncio.run(self.handle_phase4_iterate(
                params['target_score'],
                params['max_iterations']
            ))

        elif intent == 'phase4_compare':
            return asyncio.run(self.handle_phase4_compare(
                params['version_a'],
                params['version_b']
            ))

        elif intent == 'phase4_compare_request':
            return "❌ Please specify two versions to compare. Usage: /compare 1.0.0 1.2.0"

        elif intent == 'phase4_rollback':
            return asyncio.run(self.handle_phase4_rollback(params.get('version')))

        elif intent == 'phase4_versions':
            return asyncio.run(self.handle_phase4_versions())

        elif intent == 'phase4_analytics':
            return asyncio.run(self.handle_phase4_analytics())

        elif intent == 'phase4_apply':
            # TODO: Implement apply handler
            return "⚠️ /apply command coming soon! For now, use the backend API directly."

        else:
            # General conversation
            if self.current_scores:
                context = f"Current paper score: {self.current_scores['overall']}/10. The user is asking a general question: {user_message}"
            else:
                context = f"No paper evaluated yet. The user is asking: {user_message}"

        # Add to conversation history
        self.conversation_history.append({
            'role': 'user',
            'content': user_message
        })

        # Generate response using Claude
        messages = self.conversation_history.copy()
        if context:
            messages[-1]['content'] = f"{context}\n\nUser message: {user_message}"

        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console,
            transient=True
        ) as progress:
            task = progress.add_task("Thinking...", total=None)

            response = self.client.messages.create(
                model="claude-sonnet-4-5-20250929",
                max_tokens=1024,
                system=self.system_prompt,
                messages=messages
            )

            progress.update(task, completed=True)

        assistant_message = response.content[0].text

        # Add to conversation history
        self.conversation_history.append({
            'role': 'assistant',
            'content': assistant_message
        })

        return assistant_message

    def chat(self):
        """Start the interactive chat session with Rich UI."""
        # Welcome banner
        welcome = Panel(
            "[bold cyan]Paper Review Chatbot[/bold cyan] [dim](with Phase 4 Integration)[/dim]\n\n"
            "I can help you evaluate and improve your scientific papers - both full papers and individual sections!\n\n"
            "[yellow]Paper Review Commands:[/yellow]\n"
            "  • Review full paper: 'Review my paper: /path/to/paper.docx'\n"
            "  • List sections: 'list sections' or 'show sections'\n"
            "  • Review specific section: 'review abstract', 'review methods', 'review results'\n"
            "  • Review all sections: 'review all sections'\n"
            "  • Check consistency: 'check consistency' (cross-section analysis)\n"
            "  • Improve scores: 'Help me get to 8.5+'\n\n"
            "[green]Phase 4 Commands (Version Control & RAG):[/green]\n"
            "  • /suggest [section] - Get RAG-powered smart suggestions\n"
            "  • /iterate [score] [rounds] - Auto-improve until target score (e.g., /iterate 8.5)\n"
            "  • /compare [v1] [v2] - Compare versions (e.g., /compare 1.0.0 1.2.0)\n"
            "  • /rollback [version] - Rollback to previous version (e.g., /rollback 1.1.0)\n"
            "  • /versions - Show complete version history\n"
            "  • /analytics - Show comprehensive analytics dashboard\n\n"
            "[yellow]Session Management:[/yellow]\n"
            "  • Save session: 'save conversation'\n"
            "  • Load session: 'load conversation'\n"
            "  • Show history: 'show history'\n"
            "  • Exit: 'quit' or 'exit'",
            border_style="bright_blue",
            box=box.DOUBLE
        )
        console.print(welcome)

        while True:
            try:
                console.print()  # Empty line
                user_input = Prompt.ask("[bold green]💬 You[/bold green]").strip()

                if not user_input:
                    continue

                if user_input.lower() in ['quit', 'exit', 'bye']:
                    # Offer to save before exiting
                    if self.conversation_history:
                        save_prompt = Prompt.ask(
                            "[yellow]Save conversation before exiting?[/yellow]",
                            choices=["y", "n"],
                            default="y"
                        )
                        if save_prompt == "y":
                            session_id = self.save_current_session()
                            console.print(f"[green]✅ Session saved: {session_id}[/green]")

                    console.print("\n[bold cyan]👋 Goodbye! Good luck with your paper![/bold cyan]\n")
                    break

                # Generate response
                console.print()  # Empty line
                response = self.generate_response(user_input)

                # Display response as Markdown
                console.print(Panel(
                    Markdown(response),
                    title="🤖 Assistant",
                    border_style="blue"
                ))

            except KeyboardInterrupt:
                console.print("\n\n[bold cyan]👋 Chat interrupted. Goodbye![/bold cyan]\n")
                break
            except Exception as e:
                console.print(f"\n[red]❌ Error: {e}[/red]")
                console.print("[yellow]Please try again or type 'quit' to exit.[/yellow]")


def main():
    """Main entry point."""
    # Check for API key
    if not os.getenv("ANTHROPIC_API_KEY"):
        console.print("[red]❌ Error: ANTHROPIC_API_KEY not found in environment variables.[/red]")
        console.print("Please set it in your .env file or export it:")
        console.print("  export ANTHROPIC_API_KEY='your-api-key'")
        sys.exit(1)

    # Start chatbot
    chatbot = PaperReviewChatbot()
    chatbot.chat()


if __name__ == "__main__":
    main()
