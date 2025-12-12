#!/usr/bin/env python3
"""Convert grant proposal text to DOCX format for AI-CoScientist analysis."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def create_grant_docx(text_file, output_docx):
    """Convert text file to formatted DOCX."""
    # Read the text content
    text_content = Path(text_file).read_text(encoding='utf-8')

    # Create document
    doc = Document()

    # Add title
    title = doc.add_heading('K-NeuroMind: 한국형 브레인 파운데이션 모델 개발', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add subtitle
    subtitle = doc.add_paragraph('2026년도 인공지능 분야 신규 R&D 사업 제안서')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0].font
    subtitle_format.size = Pt(14)
    subtitle_format.italic = True

    doc.add_paragraph()  # Blank line

    # Process content
    lines = text_content.split('\n')
    current_section = None

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Detect headings
        if line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
            doc.add_heading(line, level=1)
        elif line.startswith('【'):  # Enhanced sections
            p = doc.add_paragraph(line)
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(11)
        elif line.startswith('< '):  # Subsections
            doc.add_heading(line, level=2)
        elif line.startswith('- ') or line.startswith('• '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('=== '):  # Major sections
            doc.add_heading(line.replace('=', '').strip(), level=1)
        else:
            doc.add_paragraph(line)

    # Save document
    doc.save(output_docx)
    print(f"✅ DOCX created: {output_docx}")
    print(f"   Paragraphs: {len(doc.paragraphs)}")
    print(f"   File size: {Path(output_docx).stat().st_size / 1024:.1f} KB")

if __name__ == '__main__':
    text_file = '/Users/jiookcha/Documents/git/AI-CoScientist/data/grant_augmented_20251019_151009.txt'
    output_docx = '/Users/jiookcha/Documents/git/AI-CoScientist/data/grant_proposal_for_analysis.docx'

    create_grant_docx(text_file, output_docx)
