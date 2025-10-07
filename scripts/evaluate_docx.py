#!/usr/bin/env python3
"""Evaluate a .docx paper using the ensemble scorer."""

import asyncio
import sys
from pathlib import Path
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent))

from src.services.paper.ensemble_scorer import EnsemblePaperScorer


async def evaluate_docx(docx_path: str):
    """Read and evaluate a .docx file."""
    print("=" * 80)
    print("PAPER QUALITY EVALUATION")
    print("=" * 80)
    print()

    # Read .docx file
    print(f"📄 Reading: {docx_path}")
    doc = Document(docx_path)

    # Extract text from all paragraphs
    full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    print(f"   Document length: {len(full_text)} characters")
    print(f"   Paragraphs: {len(doc.paragraphs)}")
    print()

    # Initialize scorer
    print("🔧 Initializing ensemble scorer...")
    scorer = EnsemblePaperScorer()
    print()

    # Evaluate
    print("🔍 Evaluating paper quality...")
    result = await scorer.score_paper(full_text, return_individual=True)
    print()

    # Display results
    print("=" * 80)
    print("EVALUATION RESULTS")
    print("=" * 80)
    print()

    print(f"📊 Overall Quality: {result['overall']:.2f}/10")
    print(f"🎯 Confidence: {result['confidence']:.2f}")
    print()

    print("📈 Dimensional Scores:")
    dimensions = result.get('dimensions', {})

    # Check if dimensions exist
    if dimensions:
        for dim_name in ['novelty', 'methodology', 'clarity', 'significance']:
            score = dimensions.get(dim_name, 0.0)

            # Rating
            if score >= 8.0:
                rating = "✅ Excellent"
            elif score >= 7.0:
                rating = "✅ Good"
            elif score >= 6.0:
                rating = "⚠️  Fair"
            elif score >= 5.0:
                rating = "⚠️  Needs Improvement"
            else:
                rating = "🔴 Weak"

            print(f"  {dim_name.capitalize():15s}: {score:.2f}/10  {rating}")
    else:
        print("  ⚠️  Multi-dimensional scores not available")
    print()

    # Model agreement
    individual = result.get('individual_scores', {})
    agreement = result.get('agreement', {})

    if individual:
        print("🤝 Model Scores:")
        if individual.get('gpt4') is not None:
            print(f"  GPT-4 Model:      {individual['gpt4']:.2f}/10")
        if individual.get('hybrid') is not None:
            print(f"  Hybrid Model:     {individual['hybrid']:.2f}/10")
        if individual.get('multitask') is not None:
            print(f"  Multi-task Model: {individual['multitask']:.2f}/10")

        if agreement:
            print(f"\n  Agreement Analysis:")
            print(f"    Max Difference: {agreement.get('max_difference', 0.0):.2f}")
            print(f"    Std Deviation:  {agreement.get('std_deviation', 0.0):.2f}")
            print(f"    Interpretation: {agreement.get('interpretation', 'N/A')}")
        print()

    # Recommendations
    print("💡 Improvement Recommendations:")

    # Find weakest dimensions (if available)
    if dimensions:
        dim_scores = {
            'novelty': dimensions.get('novelty', 0.0),
            'methodology': dimensions.get('methodology', 0.0),
            'clarity': dimensions.get('clarity', 0.0),
            'significance': dimensions.get('significance', 0.0)
        }

        sorted_dims = sorted(dim_scores.items(), key=lambda x: x[1])

        for dim, score in sorted_dims[:2]:  # Show 2 weakest dimensions
            if score < 7.0:
                if dim == 'novelty':
                    print(f"  • Novelty ({score:.1f}/10): Strengthen innovative aspects and unique contributions")
                elif dim == 'methodology':
                    print(f"  • Methodology ({score:.1f}/10): Improve research methods, rigor, and validation")
                elif dim == 'clarity':
                    print(f"  • Clarity ({score:.1f}/10): Enhance writing clarity, structure, and presentation")
                elif dim == 'significance':
                    print(f"  • Significance ({score:.1f}/10): Emphasize impact and importance of findings")

    if result['overall'] >= 8.0:
        print("  ✅ Paper quality is excellent!")
    elif result['overall'] >= 7.0:
        print("  ✅ Paper quality is good, minor improvements suggested")
    else:
        print("  ⚠️  Consider revising based on dimensional feedback above")

    print()
    print("=" * 80)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python evaluate_docx.py <path_to_docx>")
        sys.exit(1)

    docx_path = sys.argv[1]
    asyncio.run(evaluate_docx(docx_path))
