#!/usr/bin/env python3
"""Ultimate abstract improvement system with GPT-5."""

import sys
sys.path.insert(0, 'scripts')

from openai import OpenAI
from section_evaluator import SectionEvaluator, SectionType, EvaluationContext
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import json
import numpy as np

print('='*80)
print('GPT-5 ABSTRACT IMPROVEMENT SYSTEM')
print('Using OpenAI GPT-5 - The Most Powerful LLM')
print('='*80)

# Configuration
TARGET_SCORE = 8.5
MAX_ITERATIONS = 5
N_CANDIDATES = 3
N_EVALUATIONS = 2

# Weighted scoring for abstracts
ABSTRACT_WEIGHTS = {
    'clarity': 0.20,
    'completeness': 0.20,
    'conciseness': 0.30,
    'impact': 0.30
}

# Load API keys
anthropic_key = None
openai_key = None

with open('.env', 'r') as f:
    for line in f:
        if line.startswith('ANTHROPIC_API_KEY='):
            anthropic_key = line.strip().split('=', 1)[1]
        elif line.startswith('OPENAI_API_KEY='):
            openai_key = line.strip().split('=', 1)[1]

if not openai_key:
    print('❌ Error: OPENAI_API_KEY not found in .env file')
    print('Please add: OPENAI_API_KEY=your_key_here')
    sys.exit(1)

# Read original abstract
with open('abstract_original.txt', 'r') as f:
    original_text = f.read()

print(f'\n📄 Original Abstract: {len(original_text.split())} words')
print(f'\n⚙️  System Configuration:')
print(f'   • Model: GPT-5 (OpenAI)')
print(f'   • Evaluation: Claude Sonnet 4.5 (Anthropic)')
print(f'   • Target Score: {TARGET_SCORE}+')
print(f'   • Max Iterations: {MAX_ITERATIONS}')
print(f'   • Candidates per Iteration: {N_CANDIDATES}')
print(f'   • Evaluations per Candidate: {N_EVALUATIONS}')

# Initialize clients
evaluator = SectionEvaluator()  # Uses Anthropic for evaluation
openai_client = OpenAI(api_key=openai_key)

def calculate_weighted_score(evaluation):
    """Calculate weighted score based on abstract priorities."""
    weighted = sum(
        evaluation[dim]['score'] * ABSTRACT_WEIGHTS[dim]
        for dim in ABSTRACT_WEIGHTS
    )
    return weighted

def reliable_evaluate(text, n_evals=N_EVALUATIONS):
    """Evaluate multiple times for consistency."""
    evaluations = []
    for i in range(n_evals):
        eval_result = evaluator.evaluate_section(
            SectionType.ABSTRACT,
            text,
            EvaluationContext.STANDALONE
        )
        evaluations.append(eval_result)

    # Average scores
    averaged = {}
    for dim in ['clarity', 'completeness', 'conciseness', 'impact']:
        scores = [e[dim]['score'] for e in evaluations]
        averaged[dim] = {
            'score': np.mean(scores),
            'std': np.std(scores),
            'justification': evaluations[0][dim]['justification']
        }

    averaged['overall'] = {
        'score': np.mean([e['overall']['score'] for e in evaluations]),
        'std': np.std([e['overall']['score'] for e in evaluations]),
        'justification': evaluations[0]['overall']['justification']
    }

    averaged['strengths'] = evaluations[0].get('strengths', [])
    averaged['weaknesses'] = evaluations[0].get('weaknesses', [])
    averaged['suggestions'] = evaluations[0].get('suggestions', [])
    averaged['weighted_score'] = calculate_weighted_score(averaged)

    return averaged

def improve_with_gpt5(prompt, temperature=None):
    """Use GPT-5 for text improvement. Note: GPT-5 only supports temperature=1."""
    response = openai_client.chat.completions.create(
        model="gpt-5",
        messages=[
            {"role": "system", "content": "You are an expert scientific writing consultant specializing in abstract improvement."},
            {"role": "user", "content": prompt}
        ],
        max_completion_tokens=4096  # GPT-5 requires larger token limit
        # GPT-5 only supports temperature=1 (default), so we don't set it
    )
    return response.choices[0].message.content.strip()

# Multi-candidate strategies
strategies = [
    {
        'name': 'Ultra-Concise',
        'target_words': '150-160',
        'temperature': 0.2,
        'focus': 'EXTREME conciseness - remove every non-essential word, use shortest possible expressions'
    },
    {
        'name': 'Maximum-Clarity',
        'target_words': '165-175',
        'temperature': 0.2,
        'focus': 'SUPREME clarity - simplify all technical terms, improve accessibility dramatically'
    },
    {
        'name': 'Balanced-Elite',
        'target_words': '155-165',
        'temperature': 0.2,
        'focus': 'Perfect balance - optimize both clarity and conciseness equally with surgical precision'
    }
]

# Track improvement history
history = []
current_text = original_text

# Initial evaluation
print(f'\n🔍 Initial Evaluation (Claude Sonnet 4.5, averaging {N_EVALUATIONS} evals)...')
initial_eval = reliable_evaluate(current_text)
initial_score = initial_eval['overall']['score']
initial_weighted = initial_eval['weighted_score']

print(f'   Standard Score: {initial_score:.2f}/10 (±{initial_eval["overall"]["std"]:.2f})')
print(f'   Weighted Score: {initial_weighted:.2f}/10')

history.append({
    'iteration': 0,
    'model': 'Original',
    'text': current_text,
    'evaluation': initial_eval,
    'word_count': len(current_text.split()),
    'standard_score': initial_score,
    'weighted_score': initial_weighted
})

# Iterative improvement with GPT-5
for iteration in range(1, MAX_ITERATIONS + 1):
    print(f'\n{"="*80}')
    print(f'ITERATION {iteration} - GPT-5 Powered')
    print(f'{"="*80}')

    current_eval = history[-1]['evaluation']
    current_score = history[-1]['weighted_score']

    if current_score >= TARGET_SCORE:
        print(f'✅ Target {TARGET_SCORE} reached! (Weighted: {current_score:.2f})')
        break

    print(f'\n📊 Current Status:')
    print(f'   Weighted Score: {current_score:.2f}/10')
    print(f'   Standard Score: {history[-1]["standard_score"]:.2f}/10')
    print(f'   Word Count: {history[-1]["word_count"]}')

    # Identify weak dimensions
    weak_dimensions = []
    for dim in ['clarity', 'completeness', 'conciseness', 'impact']:
        if current_eval[dim]['score'] < 8.0:
            weak_dimensions.append({
                'name': dim,
                'score': current_eval[dim]['score'],
                'weight': ABSTRACT_WEIGHTS[dim],
                'weighted_impact': (8.0 - current_eval[dim]['score']) * ABSTRACT_WEIGHTS[dim]
            })

    weak_dimensions.sort(key=lambda x: x['weighted_impact'], reverse=True)

    if weak_dimensions:
        print(f'\n🎯 Priority Improvements:')
        for dim in weak_dimensions[:3]:
            print(f'   • {dim["name"].title()}: {dim["score"]:.2f}/10 ' +
                  f'(weight: {dim["weight"]:.0%}, impact: {dim["weighted_impact"]:.2f})')

    # Generate multiple candidates with GPT-5
    candidates = []

    for i, strategy in enumerate(strategies, 1):
        print(f'\n🔧 GPT-5 Candidate {i}: {strategy["name"]}')

        # Build comprehensive improvement prompt
        prompt = f'''Scientific writing expert: Create an elite-quality abstract using the {strategy['name']} strategy.

CURRENT VERSION (weighted score: {current_score:.2f}/10, words: {len(current_text.split())}):
{current_text}

DIMENSIONAL ANALYSIS:
'''

        for dim in ['clarity', 'completeness', 'conciseness', 'impact']:
            score = current_eval[dim]['score']
            weight = ABSTRACT_WEIGHTS[dim]
            status = '✓' if score >= 8.0 else '⚠️'
            prompt += f'{status} {dim.title()}: {score:.2f}/10 (weight: {weight:.0%})\n'

        if current_eval.get('weaknesses'):
            prompt += f'\nCRITICAL WEAKNESSES:\n'
            for j, weakness in enumerate(current_eval['weaknesses'][:4], 1):
                prompt += f'{j}. {weakness}\n'

        if current_eval.get('suggestions'):
            prompt += f'\nACTIONABLE SUGGESTIONS:\n'
            for j, suggestion in enumerate(current_eval['suggestions'][:5], 1):
                prompt += f'{j}. {suggestion}\n'

        prompt += f'''
STRATEGY: {strategy['name']}
FOCUS: {strategy['focus']}
TARGET WORDS: {strategy['target_words']}
TARGET WEIGHTED SCORE: {TARGET_SCORE}+

REQUIREMENTS:
1. Address ALL weaknesses listed above
2. Implement ALL suggestions
3. Achieve word count target: {strategy['target_words']}
4. Maximize weighted score (Conciseness & Impact are 30% each)
5. Maintain scientific rigor and completeness

Provide ONLY the improved abstract text (no explanation or commentary).'''

        # Generate with GPT-5
        candidate_text = improve_with_gpt5(prompt, strategy['temperature'])

        # Evaluate with Claude
        print(f'   📊 Evaluating (Claude, averaging {N_EVALUATIONS} evals)...')
        candidate_eval = reliable_evaluate(candidate_text)
        candidate_standard = candidate_eval['overall']['score']
        candidate_weighted = candidate_eval['weighted_score']

        print(f'   ✓ Standard: {candidate_standard:.2f}/10 (±{candidate_eval["overall"]["std"]:.2f})')
        print(f'   ✓ Weighted: {candidate_weighted:.2f}/10 ({candidate_weighted - current_score:+.2f})')
        print(f'   ✓ Words: {len(candidate_text.split())} ({len(candidate_text.split()) - len(current_text.split()):+d})')
        print(f'   ✓ Scores: C={candidate_eval["clarity"]["score"]:.1f} ' +
              f'Co={candidate_eval["completeness"]["score"]:.1f} ' +
              f'Cn={candidate_eval["conciseness"]["score"]:.1f} ' +
              f'I={candidate_eval["impact"]["score"]:.1f}')

        candidates.append({
            'strategy': strategy['name'],
            'text': candidate_text,
            'evaluation': candidate_eval,
            'standard_score': candidate_standard,
            'weighted_score': candidate_weighted,
            'word_count': len(candidate_text.split())
        })

    # Select best candidate
    best_candidate = max(candidates, key=lambda x: x['weighted_score'])

    print(f'\n🏆 Best GPT-5 Candidate: {best_candidate["strategy"]}')
    print(f'   Weighted: {best_candidate["weighted_score"]:.2f}/10 ({best_candidate["weighted_score"] - current_score:+.2f})')
    print(f'   Standard: {best_candidate["standard_score"]:.2f}/10')

    # Check for improvement
    if best_candidate['weighted_score'] <= current_score + 0.05:
        print(f'\n⚠️  Minimal improvement. Stopping iteration.')
        break

    # Update history
    history.append({
        'iteration': iteration,
        'model': 'GPT-5',
        'text': best_candidate['text'],
        'evaluation': best_candidate['evaluation'],
        'word_count': best_candidate['word_count'],
        'standard_score': best_candidate['standard_score'],
        'weighted_score': best_candidate['weighted_score'],
        'strategy': best_candidate['strategy']
    })

    current_text = best_candidate['text']

# Find best version
best_version = max(history, key=lambda x: x['weighted_score'])
best_iteration = best_version['iteration']
final_weighted = best_version['weighted_score']
final_standard = best_version['standard_score']

print(f'\n{"="*80}')
print(f'GPT-5 IMPROVEMENT COMPLETE')
print(f'{"="*80}')

print(f'\n🏆 Best Version: Iteration {best_iteration}')
if best_iteration > 0:
    print(f'   Model: {best_version["model"]}')
    print(f'   Strategy: {best_version["strategy"]}')

print(f'\n📊 Score Improvement:')
print(f'   Weighted: {initial_weighted:.2f} → {final_weighted:.2f} (+{final_weighted - initial_weighted:.2f})')
print(f'   Standard: {initial_score:.2f} → {final_standard:.2f} (+{final_standard - initial_score:.2f})')
print(f'   Word Count: {len(original_text.split())} → {best_version["word_count"]} ({best_version["word_count"] - len(original_text.split()):+d})')

# Save results
with open('abstract_improved_gpt5.txt', 'w') as f:
    f.write(best_version['text'])

with open('abstract_evaluation_gpt5.json', 'w') as f:
    json.dump(best_version['evaluation'], f, indent=2, ensure_ascii=False)

with open('improvement_history_gpt5.json', 'w') as f:
    history_export = []
    for h in history:
        export_item = {
            'iteration': h['iteration'],
            'model': h['model'],
            'standard_score': h['standard_score'],
            'weighted_score': h['weighted_score'],
            'word_count': h['word_count'],
            'scores_by_dimension': {
                dim: h['evaluation'][dim]['score']
                for dim in ['clarity', 'completeness', 'conciseness', 'impact']
            }
        }
        if h['iteration'] > 0:
            export_item['strategy'] = h['strategy']
        history_export.append(export_item)
    json.dump(history_export, f, indent=2)

print(f'\n✅ Files Saved:')
print(f'   • abstract_improved_gpt5.txt')
print(f'   • abstract_evaluation_gpt5.json')
print(f'   • improvement_history_gpt5.json')

# Generate comprehensive report
print(f'\n📊 Generating Comprehensive Report...')

doc = Document()

title = doc.add_heading('GPT-5 Abstract Improvement Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(11)
subtitle_run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# Executive Summary
doc.add_heading('Executive Summary', 1)

summary = doc.add_paragraph()
summary.add_run('Improvement Model: ').bold = True
summary.add_run('GPT-5 (OpenAI) - The Most Powerful LLM')

eval_model = doc.add_paragraph()
eval_model.add_run('Evaluation Model: ').bold = True
eval_model.add_run('Claude Sonnet 4.5 (Anthropic)')

doc.add_paragraph()

iterations_para = doc.add_paragraph()
iterations_para.add_run('Iterations: ').bold = True
iterations_para.add_run(f'{len(history) - 1} ({N_CANDIDATES} candidates each)')

score_para = doc.add_paragraph()
score_para.add_run('Weighted Score: ').bold = True
score_para.add_run(f'{initial_weighted:.2f} → {final_weighted:.2f} ')
improvement_run = score_para.add_run(f'(+{final_weighted - initial_weighted:.2f})')
improvement_run.bold = True
if final_weighted - initial_weighted > 0:
    improvement_run.font.color.rgb = RGBColor(0, 128, 0)

standard_para = doc.add_paragraph()
standard_para.add_run('Standard Score: ').bold = True
standard_para.add_run(f'{initial_score:.2f} → {final_standard:.2f} (+{final_standard - initial_score:.2f})')

word_para = doc.add_paragraph()
word_para.add_run('Word Count: ').bold = True
word_para.add_run(f'{len(original_text.split())} → {best_version["word_count"]} ({best_version["word_count"] - len(original_text.split()):+d})')

doc.add_page_break()

# Improvement Progression
doc.add_heading('Improvement Progression', 1)

if len(history) > 1:
    prog_table = doc.add_table(rows=len(history) + 1, cols=8)
    prog_table.style = 'Light Grid Accent 1'

    header_cells = prog_table.rows[0].cells
    headers = ['Ver', 'Model', 'Weighted', 'Standard', 'Clarity', 'Complete', 'Concise', 'Impact']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for idx, h in enumerate(history, 1):
        row = prog_table.rows[idx]
        eval_data = h['evaluation']

        row.cells[0].text = f"V{h['iteration']}"
        row.cells[1].text = h['model'][:8]
        row.cells[2].text = f"{h['weighted_score']:.2f}"
        row.cells[3].text = f"{h['standard_score']:.2f}"
        row.cells[4].text = f"{eval_data['clarity']['score']:.2f}"
        row.cells[5].text = f"{eval_data['completeness']['score']:.2f}"
        row.cells[6].text = f"{eval_data['conciseness']['score']:.2f}"
        row.cells[7].text = f"{eval_data['impact']['score']:.2f}"

        if h['iteration'] == best_iteration:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_page_break()

# Original vs Best
doc.add_heading('Original Abstract', 1)
doc.add_paragraph(f'Words: {len(original_text.split())} | Weighted: {initial_weighted:.2f}/10')
original_para = doc.add_paragraph(original_text)
original_para.style = 'Quote'

doc.add_paragraph()

doc.add_heading(f'Best Abstract (GPT-5, Iteration {best_iteration})', 1)
if best_iteration > 0:
    doc.add_paragraph(f'Strategy: {best_version["strategy"]}')
doc.add_paragraph(f'Words: {best_version["word_count"]} | Weighted: {final_weighted:.2f}/10')
best_para = doc.add_paragraph(best_version['text'])
best_para.style = 'Intense Quote'
for run in best_para.runs:
    run.bold = True

doc.add_page_break()

# Detailed Evaluation
doc.add_heading('Detailed Evaluation (Best Version)', 1)

best_eval = best_version['evaluation']

for dim_name, dim_key in [('Clarity', 'clarity'), ('Completeness', 'completeness'),
                          ('Conciseness', 'conciseness'), ('Impact', 'impact')]:
    dim_data = best_eval[dim_key]
    weight = ABSTRACT_WEIGHTS[dim_key]

    p = doc.add_paragraph()
    p.add_run(f'{dim_name}: ').bold = True
    p.add_run(f'{dim_data["score"]:.2f}/10 (±{dim_data["std"]:.2f}) ')
    p.add_run(f'[Weight: {weight:.0%}]')

    justification = doc.add_paragraph(dim_data['justification'])
    justification.paragraph_format.left_indent = Inches(0.3)
    justification.paragraph_format.space_after = Pt(6)

doc.add_paragraph()

doc.add_heading('Strengths', 2)
for strength in best_eval.get('strengths', []):
    p = doc.add_paragraph(strength, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

doc.add_paragraph()

doc.add_heading('Remaining Areas for Improvement', 2)
for weakness in best_eval.get('weaknesses', []):
    p = doc.add_paragraph(weakness, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

output_file = 'abstract_gpt5_improvement_report.docx'
doc.save(output_file)

print(f'   ✅ {output_file}')

# Final Summary
print(f'\n{"="*80}')
print(f'GPT-5 FINAL SUMMARY')
print(f'{"="*80}')

print(f'\n📊 Score Progression:')
for h in history:
    symbol = '🏆' if h['iteration'] == best_iteration else '  '
    model_str = f" [{h['model']}]"
    strategy_str = f" - {h['strategy']}" if h['iteration'] > 0 else ''
    print(f'{symbol} V{h["iteration"]}{model_str}{strategy_str}: {h["weighted_score"]:.2f}/10 ({h["word_count"]} words)')

print(f'\n📈 Dimensional Improvements:')
for dim in ['clarity', 'completeness', 'conciseness', 'impact']:
    initial_dim = initial_eval[dim]['score']
    final_dim = best_eval[dim]['score']
    change = final_dim - initial_dim
    weight = ABSTRACT_WEIGHTS[dim]
    print(f'   {dim.title():15} {initial_dim:.2f} → {final_dim:.2f} ({change:+.2f}) [weight: {weight:.0%}]')

print(f'\n🎯 Target Achievement:')
print(f'   Target: {TARGET_SCORE}+')
print(f'   Achieved (Weighted): {final_weighted:.2f}')
if final_weighted >= TARGET_SCORE:
    print(f'   🎉 TARGET ACHIEVED WITH GPT-5!')
else:
    print(f'   Gap: {TARGET_SCORE - final_weighted:.2f}')

print(f'\n✅ GPT-5 Improvement Complete!')
