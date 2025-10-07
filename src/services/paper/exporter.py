"""Paper export service for generating Word documents."""

from uuid import UUID
from pathlib import Path
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from src.models.project import Paper, PaperSection


class PaperExporter:
    """Export papers to various formats."""

    def __init__(self, db: AsyncSession):
        """Initialize exporter.

        Args:
            db: Database session
        """
        self.db = db

    async def _get_paper_with_sections(self, paper_id: UUID) -> Paper:
        """Get paper with all sections loaded.

        Args:
            paper_id: Paper ID

        Returns:
            Paper with sections

        Raises:
            ValueError: If paper not found
        """
        query = select(Paper).where(Paper.id == paper_id)
        query = query.options(selectinload(Paper.sections))

        result = await self.db.execute(query)
        paper = result.scalar_one_or_none()

        if not paper:
            raise ValueError(f"Paper {paper_id} not found")

        return paper

    def _add_title_page(self, doc: Document, paper: Paper) -> None:
        """Add title page to document.

        Args:
            doc: Word document
            paper: Paper object
        """
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(paper.title)
        title_run.font.size = Pt(18)
        title_run.font.bold = True

        # Add spacing
        doc.add_paragraph()

        # Authors (if available)
        if hasattr(paper, 'authors') and paper.authors:
            authors = doc.add_paragraph()
            authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
            authors_run = authors.add_run(paper.authors)
            authors_run.font.size = Pt(12)

            doc.add_paragraph()

        # Metadata
        metadata = doc.add_paragraph()
        metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
        metadata_text = f"Version {paper.version}\n"
        metadata_text += f"Status: {paper.status.upper()}\n"
        metadata_text += f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        metadata_run = metadata.add_run(metadata_text)
        metadata_run.font.size = Pt(10)
        metadata_run.font.color.rgb = RGBColor(128, 128, 128)

        # Page break
        doc.add_page_break()

    def _add_abstract(self, doc: Document, paper: Paper) -> None:
        """Add abstract section.

        Args:
            doc: Word document
            paper: Paper object
        """
        if paper.abstract:
            # Abstract heading
            heading = doc.add_heading('Abstract', level=1)

            # Abstract content
            abstract_para = doc.add_paragraph(paper.abstract)
            abstract_para.paragraph_format.space_after = Pt(12)

            doc.add_paragraph()

    def _add_section(self, doc: Document, section: PaperSection) -> None:
        """Add a paper section to document.

        Args:
            doc: Word document
            section: Paper section
        """
        # Section heading
        section_title = section.name.replace('_', ' ').title()
        heading = doc.add_heading(section_title, level=1)

        # Section content - split into paragraphs
        paragraphs = section.content.split('\n\n')

        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.paragraph_format.space_after = Pt(6)
                para.paragraph_format.line_spacing = 1.5

    def _format_document(self, doc: Document) -> None:
        """Apply consistent formatting to document.

        Args:
            doc: Word document
        """
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    async def export_to_word(
        self,
        paper_id: UUID,
        output_path: Optional[str] = None,
        include_metadata: bool = True
    ) -> str:
        """Export paper to Word document.

        Args:
            paper_id: Paper ID to export
            output_path: Optional custom output path
            include_metadata: Whether to include metadata (version, status, etc.)

        Returns:
            Path to generated Word document

        Raises:
            ValueError: If paper not found or has no sections
        """
        # Get paper with sections
        paper = await self._get_paper_with_sections(paper_id)

        # Check if paper has content
        if not paper.sections and not paper.abstract:
            raise ValueError(f"Paper {paper_id} has no content to export")

        # Create document
        doc = Document()

        # Add title page if metadata included
        if include_metadata:
            self._add_title_page(doc, paper)
        else:
            # Just add title
            title = doc.add_heading(paper.title, level=0)
            doc.add_paragraph()

        # Add abstract
        self._add_abstract(doc, paper)

        # Add sections in order
        sections = sorted(paper.sections, key=lambda s: s.order)
        for section in sections:
            self._add_section(doc, section)

        # Apply formatting
        self._format_document(doc)

        # Determine output path
        if not output_path:
            # Create safe filename from title
            safe_title = "".join(c for c in paper.title if c.isalnum() or c in (' ', '-', '_')).strip()
            safe_title = safe_title.replace(' ', '_')[:50]  # Limit length
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = f"paper_{safe_title}_{timestamp}.docx"

        # Ensure .docx extension
        if not output_path.endswith('.docx'):
            output_path += '.docx'

        # Save document
        doc.save(output_path)

        return output_path

    async def export_improved_paper(
        self,
        paper_id: UUID,
        improvements: dict,
        output_path: Optional[str] = None
    ) -> str:
        """Export paper with improvements applied.

        Args:
            paper_id: Paper ID
            improvements: Dict mapping section names to improved content
            output_path: Optional output path

        Returns:
            Path to generated Word document
        """
        # Get paper
        paper = await self._get_paper_with_sections(paper_id)

        # Create document
        doc = Document()

        # Add title page
        self._add_title_page(doc, paper)

        # Add note about improvements
        note = doc.add_paragraph()
        note_run = note.add_run("Note: This document contains AI-improved content")
        note_run.font.italic = True
        note_run.font.color.rgb = RGBColor(0, 0, 255)
        doc.add_paragraph()
        doc.add_page_break()

        # Add abstract
        self._add_abstract(doc, paper)

        # Add sections with improvements
        sections = sorted(paper.sections, key=lambda s: s.order)
        for section in sections:
            # Use improved content if available
            if section.name in improvements:
                # Create modified section
                improved_section = PaperSection(
                    name=section.name,
                    content=improvements[section.name],
                    order=section.order
                )
                self._add_section(doc, improved_section)

                # Add note
                note = doc.add_paragraph()
                note_run = note.add_run("(Content improved by AI)")
                note_run.font.italic = True
                note_run.font.size = Pt(10)
                note_run.font.color.rgb = RGBColor(0, 128, 0)
            else:
                self._add_section(doc, section)

            doc.add_paragraph()

        # Apply formatting
        self._format_document(doc)

        # Determine output path
        if not output_path:
            safe_title = "".join(c for c in paper.title if c.isalnum() or c in (' ', '-', '_')).strip()
            safe_title = safe_title.replace(' ', '_')[:50]
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = f"paper_improved_{safe_title}_{timestamp}.docx"

        if not output_path.endswith('.docx'):
            output_path += '.docx'

        doc.save(output_path)

        return output_path
