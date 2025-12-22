#!/usr/bin/env python3
"""
Create High-Quality Word Document from Markdown
===============================================
Generates a professional .docx file from the final proposal markdown,
including all high-resolution figures.
"""

import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Setup paths
SCRIPT_DIR = Path(__file__).parent
DATA_DIR = SCRIPT_DIR
MARKDOWN_FILE = DATA_DIR / "NRF_중견연구_제안서_통합본_v6_Final_Integration.md"
OUTPUT_FILE = DATA_DIR / "NRF_중견연구_제안서_최종본.docx"
FIGURE_DIR = DATA_DIR / "templates/figures"

def clean_text(text):
    """Clean markdown syntax for plain text insertion."""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
    text = re.sub(r'\*(.*?)\*', r'\1', text)      # Italic
    text = re.sub(r'`(.*?)`', r'\1', text)        # Code
    return text

def add_markdown_to_doc(doc, md_content):
    """Parse markdown and add to Word document with styling."""
    
    lines = md_content.split('\n')
    
    # Fonts
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'  # Or 'Arial' if Korean font issue
    font.size = Pt(11)
    
    # Process line by line
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        # 1. Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
            
        # 2. Images
        elif line.startswith('![') and '](' in line:
            # Extract image path: ![Alt](path)
            match = re.search(r'\!\[(.*?)\]\((.*?)\)', line)
            if match:
                caption = match.group(1)
                img_path_str = match.group(2)
                
                # Fix relative path
                img_name = Path(img_path_str).name
                full_img_path = FIGURE_DIR / img_name
                
                if full_img_path.exists():
                    try:
                        doc.add_picture(str(full_img_path), width=Inches(6.0))
                        # Add caption
                        last_paragraph = doc.paragraphs[-1] 
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        caption_p = doc.add_paragraph(caption)
                        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_p.style = 'Caption'
                        print(f"   ✅ Inserted image: {img_name}")
                    except Exception as e:
                        print(f"   ⚠️ Failed to insert image {img_name}: {e}")
                else:
                    print(f"   ⚠️ Image not found: {full_img_path}")
        
        # 3. Lists
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(clean_text(line[2:]), style='List Bullet')
            
        # 4. Tables (Simple parsing)
        elif line.startswith('|'):
            # Detect table start
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1 # Backtrack one step
            
            if len(table_lines) >= 3: # Header, Separator, Data
                # Create table
                headers = [c.strip() for c in table_lines[0].split('|') if c.strip()]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                
                # Header row
                hdr_cells = table.rows[0].cells
                for j, header in enumerate(headers):
                    hdr_cells[j].text = clean_text(header)
                    # Bold header
                    for run in hdr_cells[j].paragraphs[0].runs:
                        run.font.bold = True
                        
                # Data rows
                for row_line in table_lines[2:]: # Skip separator
                    cols = [c.strip() for c in row_line.split('|') if c.strip()]
                    if len(cols) == len(headers):
                        row_cells = table.add_row().cells
                        for j, col in enumerate(cols):
                            row_cells[j].text = clean_text(col)
            
        # 5. Blockquotes
        elif line.startswith('> '):
            p = doc.add_paragraph(clean_text(line[2:]))
            p.style = 'Intense Quote'
            
        # 6. Horizontal Rule
        elif line.startswith('---'):
            doc.add_page_break()
            
        # 7. Normal Text
        else:
            # Bold processing within text
            parts = re.split(r'(\*\*.*?\*\*)', line)
            p = doc.add_paragraph()
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    p.add_run(part)
                    
        i += 1

def main():
    print("=" * 60)
    print("📄 Creating Professional Word Document")
    print("=" * 60)
    
    if not MARKDOWN_FILE.exists():
        print(f"❌ Markdown file not found: {MARKDOWN_FILE}")
        return

    # Initialize document
    doc = Document()
    
    # Read markdown
    with open(MARKDOWN_FILE, 'r', encoding='utf-8') as f:
        content = f.read()
        
    # Convert
    print("   Processing markdown content...")
    add_markdown_to_doc(doc, content)
    
    # Save
    doc.save(OUTPUT_FILE)
    
    print("\n" + "=" * 60)
    print("✅ Document Created Successfully!")
    print(f"💾 Saved to: {OUTPUT_FILE}")
    print("=" * 60)

if __name__ == "__main__":
    main()


